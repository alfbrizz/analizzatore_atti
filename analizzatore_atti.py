#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Pseudonimizzatore avanzato per atti legali italiani. (v7)

Novità della v7:
- Lettura dei documenti Word (.docx), tabelle comprese (pip install python-docx).
- Nuovi dati riconosciuti: importi in euro, conti correnti, carte di credito, CAP.
- Esclusione di categorie (--escludi) e dati catastali fittizi coerenti
  (--catasto-fittizio, reversibili) introdotti nella v7.

Novità della v6:
- Partita IVA riconosciuta in più forme (P. IVA, P.I.V.A., Cod. IVA, VAT,
  prefisso IT...) e società con spaziatura irregolare (S. R. L., colonne PDF).
- Alias/denominazioni: quando un soggetto è introdotto con un nome breve
  ("di seguito «Pippo»", "per comodità X"), tutte le occorrenze successive di
  quel nome vengono pseudonimizzate.
- Revisione interattiva dei termini dubbi (--revisione, menu a testo e GUI):
  l'utente decide caso per caso; le scelte sono ricordate in un file globale
  (~/.pseudonimizzatore_atti/decisioni.json) e riapplicate in tutte le esecuzioni.

Novità della v5:
- Modalità interattiva guidata: avviando senza argomenti (o con -i) parte un
  menu a domande; ora spiega file vs cartella con esempi e accetta il
  trascinamento del percorso.
- Avvio con un clic: pausa finale in modalità interattiva così, aperto con un
  doppio clic, il programma non chiude la finestra prima di mostrare i risultati.
- OCR per atti scansionati e immagini: i PDF senza testo vengono letti pagina
  per pagina via OCR; accettati anche file immagine (.png .jpg .tiff...);
  nuova opzione --ocr per forzare l'OCR; usa la lingua italiana se disponibile.

Correzioni rispetto alla v3 (introdotte nella v4):
- TOPONIMO: la prima parola non può più essere una preposizione nuda, quindi
  "in corso di causa", "in via di definizione" ecc. non sono più falsi
  positivi INDIRIZZO ("via di Ripetta" e simili restano coperti).
- LUOGO_NASCITA: lookahead esteso (a-capo, ":", congiunzioni "e"/"ed"),
  prima non matchava a fine riga né in "nata a Milano e residente...".
- Audit: usa TUTTI i pattern di una categoria (prima il secondo pattern
  TELEFONO, cellulari senza etichetta, era escluso); nuova categoria
  LUOGO_NASCITA_RESIDUO; INDIRIZZO_RESIDUO senza falsi positivi
  ("via preliminare"); max_esempi rispettato davvero.
- leggi_txt: cp1252 provato PRIMA di latin-1 (latin-1 non fallisce mai,
  quindi il ramo cp1252 era irraggiungibile e i file Windows con ’ e –
  venivano decodificati male).
- elabora: un file che fallisce (es. PDF scansionato) non abortisce più
  l'intero fascicolo prima del salvataggio della mappa; la password di
  --encrypt-map è chiesta e validata PRIMA dell'elaborazione.
- _residui_da_verificare.txt contiene dati reali: ora è scritto 0600 come
  mappa e report, e l'avviso finale lo menziona.
- Sostituzione su gruppo basata sugli span del match (str.replace poteva
  colpire un'occorrenza identica precedente dentro il match).
- Range Unicode corretti (esclusi × U+00D7 e ÷ U+00F7); minimo di
  iterazioni KDF quando si decifra una mappa.

Crediti: realizzato dall'autore con il supporto dell'Intelligenza Artificiale (AI).
Un ringraziamento agli amici Carlo e Danilo per i suggerimenti e i test.
"""

from __future__ import annotations
import argparse, base64, getpass, json, os, re, stat, sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Dict, List, Optional, Tuple

try:
    import pypdf  # type: ignore
except ImportError:
    pypdf = None

try:
    import docx as _python_docx  # type: ignore  # pacchetto: python-docx
except ImportError:
    _python_docx = None

_NLP = None
try:
    import spacy  # type: ignore
    for _MODELLO in ("it_core_news_lg", "it_core_news_md", "it_core_news_sm"):
        try:
            _NLP = spacy.load(_MODELLO); break
        except Exception:
            continue
except ImportError:
    pass

# OCR opzionale: pytesseract + Pillow per riconoscere il testo dalle immagini,
# PyMuPDF (fitz) per rasterizzare le pagine dei PDF scansionati.
try:
    import pytesseract  # type: ignore
    from PIL import Image  # type: ignore
except ImportError:
    pytesseract = None
    Image = None
try:
    import fitz  # type: ignore  # PyMuPDF
except ImportError:
    fitz = None

_ESTENSIONI_IMG = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".gif", ".webp"}


def _tesseract_ok() -> bool:
    if pytesseract is None:
        return False
    try:
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def _lingua_ocr() -> Optional[str]:
    """Usa l'italiano se il pacchetto lingua è presente, altrimenti il default."""
    try:
        if "ita" in pytesseract.get_languages(config=""):
            return "ita"
    except Exception:
        pass
    return None


def ocr_disponibile_immagini() -> bool:
    return _tesseract_ok() and Image is not None


def ocr_disponibile_pdf() -> bool:
    return ocr_disponibile_immagini() and fitz is not None


def _ocr_immagine(img) -> str:
    lang = _lingua_ocr()
    return pytesseract.image_to_string(img, **({"lang": lang} if lang else {}))


def _modulo_presente(nome: str) -> bool:
    try:
        __import__(nome)
        return True
    except Exception:
        return False


def stato_dipendenze() -> List[Tuple[str, str, bool, str, str]]:
    """Elenco (nome, gruppo, disponibile, descrizione, comando) delle dipendenze.
    gruppo: 'pdf' = utile per i PDF, 'opz' = facoltativa."""
    voci: List[Tuple[str, str, bool, str, str]] = []
    voci.append(("pypdf", "pdf", pypdf is not None,
                 "Lettura dei file PDF (con testo).", "pip install pypdf"))
    voci.append(("python-docx", "pdf", _python_docx is not None,
                 "Lettura dei documenti Word (.docx).", "pip install python-docx"))

    voci.append(("spaCy + modello italiano", "opz", _NLP is not None,
                 "Riconoscimento dei nomi più preciso.",
                 "pip install spacy  e poi  python -m spacy download it_core_news_lg"))

    voci.append(("cryptography", "opz", _modulo_presente("cryptography"),
                 "Cifratura della mappa con password (--encrypt-map).",
                 "pip install cryptography"))

    ocr_ok = ocr_disponibile_pdf()
    libs_ocr = (pytesseract is not None and Image is not None and fitz is not None)
    if libs_ocr and not _tesseract_ok():
        descr_ocr = "OCR per PDF scansionati e immagini (librerie OK, manca il motore Tesseract)."
    elif ocr_ok and _lingua_ocr() != "ita":
        descr_ocr = "OCR per PDF scansionati e immagini (attivo, ma senza la lingua italiana)."
    else:
        descr_ocr = "OCR per PDF scansionati e immagini."
    voci.append(("OCR: pytesseract + pillow + pymupdf + motore Tesseract", "opz", ocr_ok,
                 descr_ocr,
                 "pip install pytesseract pymupdf pillow  +  motore Tesseract con lingua 'ita'"))
    return voci


def stampa_stato_dipendenze() -> None:
    voci = stato_dipendenze()
    def riga(nome, ok, descr, cmd):
        print(f"  {'[ OK ]' if ok else '[MANCA]'}  {nome}")
        print(f"          {descr}")
        if not ok:
            print(f"          -> {cmd}")
    print("-" * 70)
    print("DIPENDENZE")
    print(" I file .txt funzionano sempre, senza installare nulla.")
    print("\n Utile per i PDF con testo:")
    for nome, g, ok, descr, cmd in voci:
        if g == "pdf":
            riga(nome, ok, descr, cmd)
    print("\n Facoltative (migliorano il programma):")
    for nome, g, ok, descr, cmd in voci:
        if g == "opz":
            riga(nome, ok, descr, cmd)
    print("-" * 70)


# ---- disclaimer / scarico di responsabilità ----
DISCLAIMER_TITOLO = "AVVISO IMPORTANTE — VERIFICA SEMPRE L'ATTO PSEUDONIMIZZATO"
DISCLAIMER_TESTO = (
    "Questo strumento riduce fortemente la presenza di dati sensibili, ma NON "
    "garantisce la loro rimozione totale: nessun sistema automatico è perfetto e "
    "qualche dato può sfuggire. Prima di affidare il file pseudonimizzato a "
    "un'Intelligenza Artificiale (o a chiunque altro), rileggilo SEMPRE con "
    "attenzione e controlla il file dei residui, verificando che non siano rimasti "
    "nomi, denominazioni, numeri o altri dati riconoscibili. L'uso del programma è a "
    "responsabilità dell'utente: l'autore non risponde di eventuali dati sfuggiti "
    "alla pseudonimizzazione né delle conseguenze del loro trattamento."
)

def stampa_disclaimer() -> None:
    import textwrap
    print("\n" + "!" * 70)
    print(DISCLAIMER_TITOLO)
    print("!" * 70)
    for riga in textwrap.wrap(DISCLAIMER_TESTO, 70):
        print(riga)
    print("!" * 70)


# ---- autore e crediti ----
CREDITI_AUTORE = "Realizzato dall'autore con il supporto dell'Intelligenza Artificiale (AI)."
CREDITI_RINGRAZIAMENTI = "Un ringraziamento agli amici Carlo e Danilo per i suggerimenti e i test."

def stampa_crediti() -> None:
    print("\n" + "-" * 70)
    print(CREDITI_AUTORE)
    print(CREDITI_RINGRAZIAMENTI)
    print("-" * 70)


PLACEHOLDER_RE = re.compile(r"\[[A-Z_]+_\d+\]")

# Classi di caratteri: esclusi × (U+00D7) e ÷ (U+00F7) dai range accentati.
_MAIU = "A-ZÀ-ÖØ-Ý"
_LETT = "A-Za-zÀ-ÖØ-öø-ÿ"

MESI = ("gennaio|febbraio|marzo|aprile|maggio|giugno|luglio|"
        "agosto|settembre|ottobre|novembre|dicembre")
DATA_NUM = r"\b\d{1,2}[/\.\-]\d{1,2}[/\.\-]\d{2,4}\b"
DATA_TESTUALE = rf"\b\d{{1,2}}\s+(?:{MESI})\s+\d{{4}}\b"
DATA_QUALSIASI = rf"(?:{DATA_NUM}|{DATA_TESTUALE})"

_PAROLA = rf"(?:[{_MAIU}][{_LETT}]{{1,}}|[{_MAIU}][’'][{_MAIU}][{_LETT}]{{1,}})"

# Parole che non possono far parte di un nome di persona (confronto case-insensitive).
_NON_PERSONA = {
    # --- enti, uffici, organi ---
    "tribunale","corte","cassazione","suprema","costituzionale","appello","giudice",
    "giudici","comune","provincia","regione","regionale","statale","nazionale","comunale",
    "provinciale","ministero","ministro","agenzia","entrate","dogane","repubblica","banca",
    "condominio","societa","società","srl","spa","sas","snc","cooperativa","coop","fondazione",
    "associazione","ente","enti","istituto","istituti","camera","commercio","consiglio",
    "ordine","forense","foro","collegio","sezione","sezioni","distretto","circondario",
    "procura","procuratore","questura","prefettura","comando","carabinieri","polizia",
    "guardia","finanza","inps","inail","asl","usl","azienda","sanitaria","ospedale",
    # --- ruoli processuali / parti ---
    "convenuto","convenuta","convenuti","attore","attrice","attori","ricorrente","ricorrenti",
    "resistente","resistenti","appellante","appellato","opponente","opposto","controparte",
    "parte","parti","terzo","chiamato","intervenuto","creditore","debitore","istante",
    "rappresentato","rappresentata","rappresentante","difeso","difesa","difensore","patrocinio",
    "contro","oggetto","nonche","nonché",
    # --- atti / documenti / fasi ---
    "atto","atti","citazione","ricorso","comparsa","memoria","memorie","difensiva","conclusioni",
    "conclusionale","replica","note","nota","verbale","udienza","provvedimento","sentenza",
    "ordinanza","decreto","ingiuntivo","precetto","pignoramento","perizia","consulenza",
    "relazione","tecnica","perizale","allegato","allegati","documento","documenti","fascicolo",
    "istanza","eccezione","domanda","deposito","notifica","notificazione","costituzione",
    # --- concetti / branche del diritto ---
    "codice","fiscale","partita","iva","civile","penale","amministrativo","tributario",
    "societario","fallimentare","processuale","procedura","diritto","legge","leggi","norma",
    "norme","normativa","regolamento","articolo","art","comma","commi","principio","principi",
    "clausola","clausole","contratto","contratti","locazione","comodato","compravendita",
    "commerciale","costituzione","italiana","italiano","europea","europeo","giustizia",
    "pubblico","ministero","cancelleria","cancelliere","ruolo","registro","generale",
    "prima","primo","seconda","secondo","terza","terzo","quarta","quarto","quinta","quinto",
    "ordinario","ordinaria","speciale","straordinario","unico","unica",
    # --- indirizzi / topografia / cortesia ---
    "via","piazza","viale","corso","largo","vicolo","strada","località","contrada","borgo",
    "frazione","san","santa","santo","studio","legale","ufficio","residente","domiciliato",
    "domiciliata","nato","nata","nascita","spettabile","egregio","egregia","chiarissimo",
    "illustrissimo","onorevole","gentile","signor","signora","signori",
    # --- visura/dati catastali: intestazioni ed etichette di tabella ---
    "catasto","catastale","catastali","visura","terreni","terreno","fabbricati","fabbricato",
    "immobile","immobili","intestazione","intestatario","intestatari","dati","dato","identificativi",
    "identificativo","classamento","informazioni","informazione","foglio","particella","particelle",
    "subalterno","sub","porzione","porz","qualità","qualita","classe","superficie","reddito","redditi",
    "dominicale","agrario","agraria","anagrafici","anagrafico","anagrafe","tributaria","tributario",
    "oneri","reali","derivanti","frazionamento","scrittura","privata","sede","registrazione","volume",
    "ulteriori","ulteriore","euro","mappale","planimetria","fine","ur",
}
# Preposizioni/articoli: non possono essere la PRIMA parola di un nome di persona.
_PREP_ARTICOLI = {
    "di","de","del","dello","della","dei","degli","delle","da","dal","dalla",
    "lo","la","le","il","i","gli","l","d","e","ed","ai","al","alla","con","per",
}
_NON_PERSONA_RX = r"(?i:(?!(?:" + "|".join(map(re.escape, sorted(_NON_PERSONA))) + r")\b))"
_PAROLA_PERSONA = rf"{_NON_PERSONA_RX}{_PAROLA}"
# Separatore tra le parole di un nome: UN solo spazio, oppure un trattino.
# Evita di unire colonne di tabella separate da più spazi (es. visure catastali).
_SEP_NOME = r"(?:[ ]|[ ]?-[ ]?)"
NOME_PERSONA_1_4 = rf"{_PAROLA_PERSONA}(?:{_SEP_NOME}{_PAROLA_PERSONA}){{0,3}}"
NOME_PERSONA_2_4 = rf"{_PAROLA_PERSONA}(?:{_SEP_NOME}{_PAROLA_PERSONA}){{1,3}}"

# FIX v4: un toponimo può iniziare con al più 2 preposizioni ("via di Ripetta",
# "Piazza della Repubblica") ma DEVE poi contenere una parola con la maiuscola:
# "corso di causa", "via del tutto", "via di definizione" non matchano più.
_TOPO_PREP = r"(?:della|delle|degli|dei|del|di|da|san|santa|santo|ss\.)"
_TOPO_WORD = rf"(?:[{_MAIU}][{_LETT}'’]{{1,}}|{_TOPO_PREP})"
TOPONIMO = rf"(?:{_TOPO_PREP}[ ]+){{0,2}}[{_MAIU}][{_LETT}'’]{{1,}}(?:[ \-]+{_TOPO_WORD}){{0,5}}"

TITOLI = (r"Avv(?:\.|ocato|ocata)?|Dott\.ssa|Dott|Dr|Sig\.ra|Sig\.na|Sig|Ing|Geom|Arch|Rag|"
          r"Prof\.ssa|Prof|On|Notaio|Notaia|Giudice|C\.T\.U\.|CTU|C\.T\.P\.|CTP|P\.M\.|PM|"
          r"G\.I\.|G\.U\.")

NON_PERSONA_WORDS = {w.capitalize() for w in _NON_PERSONA} | {w.upper() for w in _NON_PERSONA}

# Sigle societarie e nome che le precede (usati sia per il pattern SOCIETA sia
# per estrarre il nome "senza sigla" da pseudonimizzare nelle occorrenze successive).
_SIGLE_SOCIETA = (
    r"S\.?\s*R\.?\s*L\.?(?:\s*S\.?)?|S\.?\s*P\.?\s*A\.?|S\.?\s*N\.?\s*C\.?|S\.?\s*A\.?\s*S\.?|"
    r"S\.?\s*C\.?\s*A\.?\s*R\.?\s*L\.?|S\.?\s*C\.?\s*R\.?\s*L\.?|"
    r"Soc\.?\s*Coop\.?(?:\s*a\s*r\.?\s*l\.?)?|Società\s+Cooperativa|Cooperativa|"
    r"Associazione|Fondazione"
)
# Nota: niente '.' nei token del nome, così il match non attraversa il confine
# di frase (es. "...Cassazione. Beta Costruzioni S.p.A." non diventa un unico nome).
_NOME_SOCIETA = (rf"[{_MAIU}][{_LETT}0-9&'’\-]*"
                 rf"(?:[ \t]+[{_MAIU}0-9][{_LETT}0-9&'’\-]*){{0,5}}")
_SOCIETA_NAME_RE = re.compile(rf"\b(?P<nome>{_NOME_SOCIETA})[ \t]+(?i:{_SIGLE_SOCIETA})\b")


@dataclass(frozen=True)
class PatternDef:
    categoria: str
    pattern: str
    flags: int = 0
    gruppo: Optional[str] = None
    categoria_placeholder: Optional[str] = None


PATTERN_STRUTTURATI: List[PatternDef] = [
    PatternDef("CODICE_FISCALE",
        r"\b[A-Z]{6}[0-9LMNPQRSTUV]{2}[ABCDEHLMPRST][0-9LMNPQRSTUV]{2}"
        r"[A-Z][0-9LMNPQRSTUV]{3}[A-Z]\b", re.IGNORECASE),
    PatternDef("IBAN",
        r"\bIT\d{2}[ ]?[A-Z][ ]?(?:\d[ ]?){10}(?:[0-9A-Z][ ]?){12}\b", re.IGNORECASE),
    PatternDef("CARTA_CREDITO",   # 16 cifre in 4 gruppi (dopo IBAN, così l'IBAN è già stato tolto)
        r"\b\d{4}[ \-]\d{4}[ \-]\d{4}[ \-]\d{4}\b"),
    PatternDef("CARTA_CREDITO",   # con etichetta esplicita
        r"\b(?:carta\s+di\s+credito|carta\s+n|pan)\.?\s*(?:n\.?|nr\.?|numero)?\s*(?:\d[ \-]?){13,19}\b",
        re.IGNORECASE),
    PatternDef("CONTO_CORRENTE",
        r"\b(?:c/c|c\.\s*c\.|conto\s+corrente)\s*(?:bancario|postale)?\s*"
        r"(?:n\.?|nr\.?|numero)?\s*[\d./\-]{5,20}\b", re.IGNORECASE),
    PatternDef("PEC",
        r"\b(?:pec|posta\s+elettronica\s+certificata)[:\s]*[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b",
        re.IGNORECASE),
    PatternDef("EMAIL", r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"),
    PatternDef("PARTITA_IVA",   # con etichetta (molte varianti: P.IVA, P. I.V.A., cod. IVA, VAT...)
        r"\b(?:partita\s*i\.?\s*v\.?\s*a\.?|p\.?\s*iva|p\.?\s*i\.?\s*v\.?\s*a\.?|"
        r"cod(?:ice)?\.?\s*iva|vat(?:\s*(?:number|n\.?))?)\b[\s:n\.º°/-]*(?:IT)?\s*\d{11}\b",
        re.IGNORECASE),
    PatternDef("PARTITA_IVA",   # forma con prefisso IT (es. IT01234567890)
        r"\bIT\s?\d{11}\b"),
    PatternDef("IMPORTO",   # somme in euro (€ 1.234,56 · 1.234,56 euro · EUR 500)
        r"(?:\b(?:euro|eur)\b|€)\s*\d{1,3}(?:[.\s]\d{3})*(?:,\d{1,2})?"
        r"|\d{1,3}(?:[.\s]\d{3})*(?:,\d{1,2})?\s*(?:\b(?:euro|eur)\b|€)", re.IGNORECASE),
    PatternDef("TELEFONO",   # con etichetta esplicita
        r"\b(?:tel(?:efono)?|cell(?:ulare)?|mobile|fax|recapito)\.?\s*[:\-]?\s*"
        r"(?:\+39\s*)?(?:0\d{1,4}|3\d{2})[\s.\-/]?\d{5,8}\b", re.IGNORECASE),
    PatternDef("TELEFONO",   # cellulare senza etichetta, ancorato ai prefissi mobili
        r"(?<!\d)(?:\+39\s*)?3[2-9]\d[\s.\-/]?\d{6,7}\b", re.IGNORECASE),
    PatternDef("NUMERO_RUOLO",
        r"\bR\.?\s*G\.?\s*(?:N\.?\s*R\.?|E\.?|A\.?\s*C\.?|Trib\.?)?\s*(?:n\.?|nr\.?|numero)?\s*"
        r"\d{1,8}\s*/\s*\d{2,4}\b", re.IGNORECASE),
    PatternDef("NUMERO_PROVVEDIMENTO",
        r"\b(?:sentenza|ordinanza|decreto(?:\s+ingiuntivo)?|provvedimento|procedimento)"
        r"\s*(?:n\.?|nr\.?|numero)?\s*\d{1,8}\s*/\s*\d{2,4}\b", re.IGNORECASE),
    PatternDef("REPERTORIO_NOTARILE",
        r"\b(?:rep(?:ertorio)?\.|racc(?:olta)?\.)\s*(?:n\.?|nr\.?|numero)?\s*\d{1,10}\b",
        re.IGNORECASE),
    PatternDef("DOCUMENTO_IDENTITA",
        r"\b(?:carta\s+d[’' ]identità|carta\s+di\s+identita|c\.\s*i\.|patente|passaporto|"
        r"permesso\s+di\s+soggiorno|documento\s+d[’' ]identità)\s*"
        r"(?:n\.?|nr\.?|numero)?\s*[A-Z0-9]{5,20}\b", re.IGNORECASE),
    PatternDef("TESSERA_SANITARIA",
        r"\b(?:tessera\s+sanitaria|codice\s+tessera\s+sanitaria|t\.s\.)\s*"
        r"(?:n\.?|nr\.?|numero)?\s*[A-Z0-9]{8,25}\b", re.IGNORECASE),
    PatternDef("TARGA", r"\b[A-Z]{2}\s?\d{3}\s?[A-Z]{2}\b"),
    PatternDef("DATI_CATASTALI",
        r"\b(?:foglio|particella|mappale|subalterno|fg\.?|p\.lla|mapp\.?|sub\.?|part\.|f\.)\s*"
        r"(?:n\.?|nr\.?|numero)?\s*\d+[A-Za-z0-9/\-]*\b", re.IGNORECASE),
    PatternDef("COORDINATA",
        r"\b(?:lat(?:itudine)?|lng|long(?:itudine)?|coordinate?)\s*[:=]?\s*-?\d{1,3}[\.,]\d+\b",
        re.IGNORECASE),
    PatternDef("DATA_NASCITA",
        rf"\b(?:nato|nata|nascita)\b[^\n]{{0,80}}?\b(?:il|data)\s+(?P<data>{DATA_QUALSIASI})",
        re.IGNORECASE, gruppo="data", categoria_placeholder="DATA_NASCITA"),
    # FIX v4: lookahead esteso — copre fine riga (\r\n), ":" e le congiunzioni
    # "e"/"ed" ("nata a Milano e residente..."); il vecchio "$" senza MULTILINE
    # valeva solo a fine testo.
    PatternDef("LUOGO_NASCITA",
        rf"\b(?i:nato|nata)\s+(?i:a|in)\s+(?P<luogo>{TOPONIMO})"
        rf"(?=\s*(?:,|\(|;|\.|:|[\r\n]|$)|\s(?i:il|in|e|ed)\b)",
        0, gruppo="luogo", categoria_placeholder="LUOGO_NASCITA"),
    PatternDef("INDIRIZZO",
        rf"\b(?i:Via|V\.le|Viale|Piazza|P\.zza|Corso|C\.so|Largo|Vicolo|Strada|"
        rf"Località|Loc\.|Contrada|C\.da|Borgo|Frazione)\s+{TOPONIMO}"
        rf"(?:\s*,?\s*(?:n\.?|num\.?|numero)?\s*\d+[A-Za-z]?)?", 0),
    PatternDef("CAP",   # codice di avviamento postale, con etichetta (evita falsi positivi)
        r"\b(?:c\.a\.p\.?|cap)\b\s*[:\-]?\s*\d{5}\b", re.IGNORECASE),
    PatternDef("CONDOMINIO",
        rf"\bCondominio\s+(?:denominato\s+)?[\"“]?{TOPONIMO}[\"”]?", 0),
    # Spaziatura tollerante: il nome e la sigla possono avere più spazi (colonne
    # PDF) e la sigla può avere spazi/punti tra le lettere (S. R. L.).
    PatternDef("SOCIETA", rf"\b{_NOME_SOCIETA}[ \t]+(?i:{_SIGLE_SOCIETA})\b"),
    PatternDef("PERSONA_CON_TITOLO",
        rf"\b(?P<titolo>(?i:{TITOLI})\.?)\s+(?P<nome>{NOME_PERSONA_1_4})",
        0, gruppo="nome", categoria_placeholder="PERSONA"),
]

PATTERN_DATE_GENERICHE: List[PatternDef] = [
    PatternDef("DATA", DATA_NUM),
    PatternDef("DATA", DATA_TESTUALE, re.IGNORECASE),
]

# FIX v4: l'audit copre TUTTI i pattern di una categoria (la v3 teneva solo il
# primo: il pattern dei cellulari senza etichetta era escluso dall'audit).
_BY_CAT: Dict[str, List[PatternDef]] = {}
for _d in PATTERN_STRUTTURATI:
    _BY_CAT.setdefault(_d.categoria, []).append(_d)

def _audit_da_categoria(nome_audit: str, categoria: str) -> List[PatternDef]:
    return [PatternDef(nome_audit, d.pattern, d.flags) for d in _BY_CAT[categoria]]

AUDIT_PATTERNS: List[PatternDef] = (
    _audit_da_categoria("CODICE_FISCALE_RESIDUO", "CODICE_FISCALE")
    + _audit_da_categoria("IBAN_RESIDUO", "IBAN")
    + _audit_da_categoria("EMAIL_RESIDUA", "EMAIL")
    + _audit_da_categoria("PARTITA_IVA_RESIDUA", "PARTITA_IVA")
    + _audit_da_categoria("TELEFONO_RESIDUO", "TELEFONO")
    + _audit_da_categoria("TARGA_RESIDUA", "TARGA")
    + _audit_da_categoria("DATI_CATASTALI_RESIDUI", "DATI_CATASTALI")
    + [
        # FIX v4: usa TOPONIMO (il vecchio "\w+" IGNORECASE segnalava
        # "via preliminare", "in corso di causa" ecc.)
        PatternDef("INDIRIZZO_RESIDUO",
            rf"\b(?i:Via|Viale|Piazza|Corso|Largo|Vicolo|Strada|Località|Contrada)\s+{TOPONIMO}", 0),
        PatternDef("DATA_NASCITA_RESIDUA",
            rf"\b(?:nato|nata|nascita)\b[^\n]{{0,80}}?(?:{DATA_QUALSIASI})", re.IGNORECASE),
        # FIX v4: nuova categoria — i luoghi di nascita sfuggiti erano invisibili.
        PatternDef("LUOGO_NASCITA_RESIDUO",
            rf"\b(?i:nato|nata)\s+(?i:a|in)\s+{TOPONIMO}", 0),
    ]
)


# ---- sicurezza/mappa ----
def scrivi_file_privato(percorso: Path, contenuto: str) -> None:
    percorso.parent.mkdir(parents=True, exist_ok=True)
    fd = os.open(str(percorso), os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    with os.fdopen(fd, "w", encoding="utf-8") as f:
        f.write(contenuto)
    try:
        os.chmod(percorso, stat.S_IRUSR | stat.S_IWUSR)
    except Exception:
        pass

def _crypto():
    try:
        from cryptography.fernet import Fernet
        from cryptography.hazmat.primitives import hashes
        from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    except ImportError as exc:
        raise RuntimeError("Per --encrypt-map serve: pip install cryptography") from exc
    return Fernet, hashes, PBKDF2HMAC

def _deriva_chiave(password: str, salt: bytes, iterations: int = 480_000) -> bytes:
    _, hashes, PBKDF2HMAC = _crypto()
    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32, salt=salt, iterations=iterations)
    return base64.urlsafe_b64encode(kdf.derive(password.encode("utf-8")))

def cifra_json(dati: Dict[str, str], password: str) -> str:
    Fernet, _, _ = _crypto()
    salt = os.urandom(16); iterations = 480_000
    token = Fernet(_deriva_chiave(password, salt, iterations)).encrypt(
        json.dumps(dati, ensure_ascii=False, indent=2).encode("utf-8"))
    return json.dumps({"encrypted": True, "kdf": "PBKDF2HMAC-SHA256",
        "iterations": iterations, "salt_b64": base64.b64encode(salt).decode("ascii"),
        "payload_b64": token.decode("ascii"),
        "created_at": datetime.now().isoformat(timespec="seconds")}, ensure_ascii=False, indent=2)

def decifra_json(pkt: Dict[str, object], password: str) -> Dict[str, str]:
    Fernet, _, _ = _crypto()
    salt = base64.b64decode(str(pkt["salt_b64"]))
    # FIX v4: minimo di iterazioni (un file manomesso non può degradare il KDF).
    iterations = int(pkt.get("iterations", 480_000))
    if iterations < 100_000:
        raise ValueError("Numero di iterazioni KDF troppo basso: mappa sospetta o manomessa.")
    chiave = _deriva_chiave(password, salt, iterations)
    return json.loads(Fernet(chiave).decrypt(str(pkt["payload_b64"]).encode("ascii")).decode("utf-8"))

def chiedi_password_cifratura() -> str:
    """FIX v4: chiamata PRIMA dell'elaborazione, così un errore di password
    non fa perdere la mappa a lavoro già svolto."""
    pw = getpass.getpass("Password per cifrare la mappa: ")
    if pw != getpass.getpass("Conferma password: "):
        raise RuntimeError("Le password non coincidono. Nessun file elaborato.")
    if not pw.strip():
        raise RuntimeError("Password vuota non ammessa. Nessun file elaborato.")
    return pw

def salva_mappa(mappa: Dict[str, str], percorso: Path, password: Optional[str] = None) -> Path:
    if password is not None:
        if not str(percorso).endswith(".enc"):
            percorso = percorso.with_name(percorso.name + ".enc")
        scrivi_file_privato(percorso, cifra_json(mappa, password))
    else:
        scrivi_file_privato(percorso, json.dumps(mappa, ensure_ascii=False, indent=2))
    return percorso

def carica_mappa(percorso: Path) -> Dict[str, str]:
    dati = json.loads(percorso.read_text(encoding="utf-8"))
    if isinstance(dati, dict) and dati.get("encrypted") is True:
        return decifra_json(dati, getpass.getpass("Password per decifrare la mappa: "))
    if not isinstance(dati, dict):
        raise ValueError("Formato mappa non valido.")
    return {str(k): str(v) for k, v in dati.items()}


# ---- decisioni dell'utente (memoria persistente) ----
def _norm_termine(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip().upper()

def percorso_decisioni() -> Path:
    return Path.home() / ".pseudonimizzatore_atti" / "decisioni.json"

def carica_decisioni() -> Dict[str, object]:
    """Carica le decisioni globali dell'utente (termini da pseudonimizzare/ignorare)."""
    p = percorso_decisioni()
    try:
        if p.exists():
            d = json.loads(p.read_text(encoding="utf-8"))
            anon = {str(k): dict(v) if isinstance(v, dict) else {"categoria": "DATO_SENSIBILE", "testo": str(v)}
                    for k, v in dict(d.get("pseudonimizza", {})).items()}
            ign = [str(x) for x in d.get("ignora", [])]
            ass = []
            for g in d.get("associazioni", []):
                if isinstance(g, dict) and g.get("forme"):
                    ass.append({"categoria": str(g.get("categoria", "SOGGETTO")),
                                "forme": [str(x) for x in g["forme"] if str(x).strip()]})
            return {"pseudonimizza": anon, "ignora": ign, "associazioni": ass}
    except Exception:
        pass
    return {"pseudonimizza": {}, "ignora": [], "associazioni": []}

def salva_decisioni(dec: Dict[str, object]) -> None:
    p = percorso_decisioni()
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        scrivi_file_privato(p, json.dumps(dec, ensure_ascii=False, indent=2))
    except Exception:
        pass


# ---- associazioni per singolo atto (file affiancato, modificabile) ----
def percorso_associazioni(bersagli, cartella) -> Path:
    """File delle associazioni, con lo STESSO NOME dell'atto (o della cartella)."""
    if cartella:
        return cartella / f"{cartella.name}_associazioni.txt"
    return bersagli[0].with_name(bersagli[0].stem + "_associazioni.txt")

def assoc_path_da_percorso(percorso: str) -> Optional[Path]:
    """Ricava il file delle associazioni dal percorso indicato (per la GUI)."""
    percorso = (percorso or "").strip()
    if not percorso:
        return None
    p = Path(percorso).expanduser()
    if p.is_dir():
        return p / f"{p.name}_associazioni.txt"
    return p.with_name(p.stem + "_associazioni.txt")

def leggi_associazioni_txt(path: Path) -> List[str]:
    """Righe di associazioni (una per soggetto); ignora righe vuote e commenti."""
    righe: List[str] = []
    try:
        if path and path.exists():
            for r in path.read_text(encoding="utf-8").splitlines():
                r = r.strip()
                if r and not r.startswith("#"):
                    righe.append(r)
    except Exception:
        pass
    return righe

def scrivi_associazioni_txt(path: Path, righe: List[str], nome_atto: str = "") -> None:
    """Salva le associazioni in un file leggibile e modificabile a mano."""
    intest = [
        f"# Associazioni di nomi per: {nome_atto}".rstrip(),
        "# Una riga per ogni soggetto; forme separate da virgola.",
        "# Tutte le forme di una riga diventano lo stesso segnaposto.",
        "# Le righe che iniziano con # sono ignorate. Puoi modificare questo file.",
        "",
    ]
    contenuto = "\n".join(intest + [r.strip() for r in righe if r.strip()]) + "\n"
    try:
        scrivi_file_privato(path, contenuto)
    except Exception:
        pass


# Trigger che introducono un alias/denominazione abbreviata di una parte o società.
_ALIAS_TRIGGER = (
    r"denominat[oa]|indicat[oa]|chiamat[oa]|abbreviat[oa]\s+in|"
    r"per\s+(?:comodità|comodita|brevità|brevita)|breviter|"
    r"di\s+seguito|in\s+seguito|qui\s+di\s+seguito|nel\s+prosieguo|"
    r"d['’]ora\s+in\s+(?:avanti|poi)|sig\.?la"
)
# Parole di riempimento ammesse fra il trigger e l'alias ("di seguito ANCHE LA ...").
_ALIAS_FILLER = (
    r"(?:anche|pure|solo|breviter|la|il|lo|le|gli|una|un|detta|detto|"
    r"denominat[oa]|indicat[oa]|nominat[oa]|abbreviat[oa]|come|parte|"
    r"societ[àa]|ditta|impresa|banca|ente)"
)
# Alias tra virgolette (preferito): cattura qualunque contenuto virgolettato.
_ALIAS_QUOTED_RE = re.compile(
    rf"(?i:{_ALIAS_TRIGGER})(?:\s+{_ALIAS_FILLER}){{0,3}}\s*[:,]?\s*"
    r"[«\"“'‘](?P<alias>[^«»\"“”'‘’\n]{2,50})[»\"”'’]")
# Alias "nudo" (senza virgolette): una sequenza con iniziale maiuscola.
_ALIAS_BARE_RE = re.compile(
    rf"(?i:{_ALIAS_TRIGGER})(?:\s+{_ALIAS_FILLER}){{0,3}}\s*[:,]?\s+"
    r"(?P<alias>[A-ZÀ-Ý][\wÀ-ÿ'’&\.\-]+(?:\s+[A-ZÀ-Ý0-9][\wÀ-ÿ'’&\.\-]+){0,4})")


class Pseudonimizzatore:
    def __init__(self, pseudonimizza_date=False, fallback_nomi=True):
        self.pseudonimizza_date = pseudonimizza_date
        self.fallback_nomi = fallback_nomi
        self.registro: Dict[str, str] = {}
        self.mappa: Dict[str, str] = {}
        self.contatori: Dict[str, int] = {}
        # Termini decisi dall'utente/alias: norm -> (categoria, testo da mostrare).
        self.termini_forzati: Dict[str, Tuple[str, str]] = {}
        # Termini che l'utente ha deciso di NON pseudonimizzare (norm).
        self.termini_ignora: set = set()
        # Associazioni: gruppi di forme dello stesso soggetto -> stesso segnaposto.
        # Ogni gruppo: {"categoria": str, "forme": [str, ...]}
        self.gruppi: List[dict] = []
        # Categorie da NON pseudonimizzare (codici, es. {"DATI_CATASTALI"}).
        self.categorie_escluse: set = set()
        # Categorie da sostituire con dati FITTIZI ma coerenti (non con segnaposto).
        self.categorie_fittizie: set = set()
        self._pseudo: Dict[str, str] = {}       # "TIPO::NUM" -> numero fittizio
        self._pseudo_cont: Dict[str, int] = {}

    def _cat_esclusa(self, categoria: str) -> bool:
        return categoria in self.categorie_escluse

    # --- dati fittizi coerenti (attualmente per i dati catastali) ---
    _BASE_PSEUDO = {"FOGLIO": 51, "PARTICELLA": 201, "MAPPALE": 301, "SUBALTERNO": 11, "CATASTO": 501}

    def _tipo_catastale(self, label: str) -> str:
        l = label.lower()
        if "foglio" in l or re.search(r"\bfg\b|\bf\.", l):
            return "FOGLIO"
        if "particella" in l or "p.lla" in l or "part" in l:
            return "PARTICELLA"
        if "mappale" in l or "mapp" in l:
            return "MAPPALE"
        if "subalterno" in l or "sub" in l:
            return "SUBALTERNO"
        return "CATASTO"

    def _pseudonimo_catastale(self, testo_match: str) -> Optional[str]:
        """Sostituisce il numero catastale con uno fittizio, coerente (stesso
        numero reale -> stesso fittizio) e reversibile (salvato nella mappa)."""
        m = re.match(r"^(?P<lab>.*?)(?P<num>\d[\w/\-]*)\s*$", testo_match, re.DOTALL)
        if not m:
            return None
        lab, num = m.group("lab"), m.group("num")
        tipo = self._tipo_catastale(lab)
        chiave = f"{tipo}::{num.upper()}"
        if chiave not in self._pseudo:
            base = self._BASE_PSEUDO.get(tipo, 501)
            self._pseudo_cont[tipo] = self._pseudo_cont.get(tipo, base - 1) + 1
            self._pseudo[chiave] = str(self._pseudo_cont[tipo])
        fittizio = f"{lab}{self._pseudo[chiave]}"
        # Reversibile: la stringa fittizia rimanda al valore reale.
        self.mappa.setdefault(fittizio.strip(), testo_match.strip())
        return fittizio

    # --- associazioni (più forme dello stesso soggetto = stesso segnaposto) ---
    def aggiungi_gruppo(self, forme, categoria: str = "SOGGETTO") -> None:
        forme = [f.strip() for f in forme if f and f.strip()]
        if len(forme) < 1:
            return
        self.gruppi.append({"categoria": categoria or "SOGGETTO", "forme": forme})

    def _placeholder_per_gruppo(self, categoria: str, forme: List[str]) -> str:
        # La forma più lunga fa da valore di ripristino (nome completo).
        rappr = max(forme, key=len)
        ph = self._segnaposto(categoria, rappr)
        # Tutte le altre forme puntano allo STESSO segnaposto.
        for f in forme:
            chiave = f"{categoria}::{self._normalizza(categoria, f.strip())}"
            self.registro[chiave] = ph
        return ph

    def _applica_gruppi(self, testo: str) -> str:
        for g in self.gruppi:
            forme = [f for f in g.get("forme", []) if f and f.strip()]
            if not forme:
                continue
            cat = g.get("categoria", "SOGGETTO")
            ph = self._placeholder_per_gruppo(cat, forme)
            for f in sorted(set(forme), key=len, reverse=True):
                corpo = re.sub(r"(?:\\ )+", r"\\s+", re.escape(f.strip()))
                rx = re.compile(rf"(?<![\wÀ-ÿ'’]){corpo}(?![\wÀ-ÿ'’])", re.IGNORECASE)
                testo = self._fuori_placeholder(testo, lambda s, rx=rx, ph=ph: rx.sub(ph, s))
        return testo

    # --- termini forzati / da ignorare ---
    def aggiungi_termine_forzato(self, testo: str, categoria: str = "DATO_SENSIBILE") -> None:
        n = _norm_termine(testo)
        if len(n) < 2 or n in self.termini_ignora:
            return
        self.termini_forzati.setdefault(n, (categoria, testo.strip()))

    def aggiungi_ignora(self, testo: str) -> None:
        n = _norm_termine(testo)
        self.termini_ignora.add(n)
        self.termini_forzati.pop(n, None)

    def _ignorato(self, testo: str) -> bool:
        return _norm_termine(testo) in self.termini_ignora

    def _applica_termini_forzati(self, testo: str) -> str:
        if not self.termini_forzati:
            return testo
        # I termini più lunghi hanno la precedenza (evita match parziali:
        # "Alfa Beta Costruzioni" prima di "Alfa Beta").
        for norm, (cat, orig) in sorted(self.termini_forzati.items(),
                                        key=lambda kv: len(kv[1][1]), reverse=True):
            corpo = re.sub(r"(?:\\ )+", r"\\s+", re.escape(orig.strip()))
            rx = re.compile(rf"(?<![\wÀ-ÿ'’]){corpo}(?![\wÀ-ÿ'’])", re.IGNORECASE)
            testo = self._fuori_placeholder(
                testo, lambda s, rx=rx, cat=cat: rx.sub(lambda m: self._segnaposto(cat, m.group(0)), s))
        return testo

    def _registra_alias(self, alias: str) -> None:
        alias = alias.strip(" .,;:\"'«»“”‘’")
        parole = [p for p in re.split(r"[\s\-]+", alias) if p]
        if not parole:
            return
        # Se sono tutte parole generiche/istituzionali non è un nome proprio.
        if all(p.lower() in _NON_PERSONA or p.lower() in _PREP_ARTICOLI for p in parole):
            return
        self.aggiungi_termine_forzato(alias, "DENOMINAZIONE")

    def _cattura_alias(self, testo: str) -> None:
        """Registra come termini da pseudonimizzare gli alias/denominazioni tipo
        'di seguito «Pippo»' o 'di seguito anche la Alfa Beta', usati poi nell'atto."""
        def seg(s: str) -> str:
            for m in _ALIAS_QUOTED_RE.finditer(s):
                self._registra_alias(m.group("alias"))
            for m in _ALIAS_BARE_RE.finditer(s):
                self._registra_alias(m.group("alias"))
            return s
        self._fuori_placeholder(testo, seg)

    def _cattura_nomi_societa(self, testo: str) -> None:
        """Dal nome completo di una società (es. 'Alfa Beta S.r.l.') registra il
        nome SENZA sigla ('Alfa Beta') così viene pseudonimizzato anche quando,
        più avanti nell'atto, la società è richiamata senza la sigla."""
        _ART = re.compile(r"(?i)^(?:l['’]\s*|la\s+|il\s+|lo\s+|le\s+|i\s+|gli\s+|"
                          r"un['’]?\s+|una\s+|dell['’]?\s+|della\s+|del\s+)")
        def seg(s: str) -> str:
            for m in _SOCIETA_NAME_RE.finditer(s):
                nome = _ART.sub("", m.group("nome").strip()).strip()
                parole = [p for p in re.split(r"[\s\-]+", nome) if p]
                signif = [p for p in parole if p.lower() not in _NON_PERSONA
                          and p.lower() not in _PREP_ARTICOLI]
                if not signif:
                    continue
                # Nome troppo corto/generico se singola parola breve.
                if len(parole) == 1 and len(nome) < 4:
                    continue
                self.aggiungi_termine_forzato(nome, "SOCIETA")
            return s
        self._fuori_placeholder(testo, seg)

    def _normalizza(self, categoria: str, valore: str) -> str:
        base = re.sub(r"\s+", " ", valore).strip().upper()
        if categoria in {"CODICE_FISCALE", "IBAN", "PARTITA_IVA"}:
            base = re.sub(r"\s+", "", base)
        if categoria == "TELEFONO":
            base = re.sub(r"[^0-9+]", "", base)
        return base

    def _segnaposto(self, categoria: str, valore: str) -> str:
        # FIX v4: strip dei soli spazi (":" e "-" facevano perdere caratteri
        # del testo originale al momento del ripristino).
        valore = valore.strip(" \t\u00a0")
        chiave = f"{categoria}::{self._normalizza(categoria, valore)}"
        if chiave in self.registro:
            return self.registro[chiave]
        self.contatori[categoria] = self.contatori.get(categoria, 0) + 1
        ph = f"[{categoria}_{self.contatori[categoria]}]"
        self.registro[chiave] = ph
        self.mappa[ph] = valore
        return ph

    def _segnaposto_preserva_bordi(self, categoria: str, valore: str) -> str:
        m = re.match(r"^(?P<pre>\s*)(?P<core>.*?)(?P<post>\s*)$", valore, flags=re.DOTALL)
        if not m:
            return self._segnaposto(categoria, valore)
        return m.group("pre") + self._segnaposto(categoria, m.group("core")) + m.group("post")

    def _fuori_placeholder(self, testo: str, fn: Callable[[str], str]) -> str:
        parti, ultimo = [], 0
        for m in PLACEHOLDER_RE.finditer(testo):
            parti.append(fn(testo[ultimo:m.start()])); parti.append(m.group(0)); ultimo = m.end()
        parti.append(fn(testo[ultimo:]))
        return "".join(parti)

    def _sostituisci(self, testo: str, d: PatternDef) -> str:
        cat = d.categoria_placeholder or d.categoria
        if self._cat_esclusa(cat):
            return testo
        def seg(s: str) -> str:
            def repl(m):
                if d.gruppo:
                    # FIX v4: ricostruzione basata sugli span (str.replace poteva
                    # sostituire un'occorrenza identica precedente nel match).
                    ini, fin = m.span(d.gruppo)
                    base = m.start()
                    tutto = m.group(0)
                    return (tutto[:ini - base]
                            + self._segnaposto(cat, m.group(d.gruppo))
                            + tutto[fin - base:])
                if cat in self.categorie_fittizie:
                    ps = self._pseudonimo_catastale(m.group(0))
                    if ps is not None:
                        return ps
                return self._segnaposto_preserva_bordi(cat, m.group(0))
            return re.sub(d.pattern, repl, s, flags=d.flags)
        return self._fuori_placeholder(testo, seg)

    def _entita_persona_valida(self, v: str) -> bool:
        """Filtra i falsi positivi di spaCy (intestazioni, etichette, blocchi di testo)."""
        if self._ignorato(v):
            return False
        # Un nome di persona sta su una sola riga: gli a-capo indicano un blocco di layout.
        if "\n" in v or "\r" in v:
            return False
        parole = [p.strip(".,;:()[]{}\"'’") for p in re.split(r"[\s\-]+", v.strip()) if p.strip()]
        if not parole:
            return False
        # Contiene un termine istituzionale/giuridico/catastale: non è una persona.
        if any(p.lower() in _NON_PERSONA for p in parole):
            return False
        # Non può iniziare con preposizione/articolo.
        if parole[0].lower() in _PREP_ARTICOLI:
            return False
        # Troppe parole: probabile frammento di testo, non un nome.
        if len(parole) > 4:
            return False
        # Deve contenere almeno una parola con iniziale maiuscola (evita spazzatura).
        if not any(p[:1].isupper() for p in parole):
            return False
        return True

    def _nomi_ner(self, testo: str) -> str:
        if _NLP is None:
            return testo
        def seg(s: str) -> str:
            doc = _NLP(s); nuovo = s
            for ent in sorted(doc.ents, key=lambda e: e.start_char, reverse=True):
                if ent.label_ in {"PER", "PERSON"}:
                    v = ent.text.strip()
                    if v and not PLACEHOLDER_RE.fullmatch(v) and self._entita_persona_valida(v):
                        nuovo = nuovo[:ent.start_char] + self._segnaposto("PERSONA", v) + nuovo[ent.end_char:]
            return nuovo
        return self._fuori_placeholder(testo, seg)

    def _sembra_nome(self, valore: str, contesto: str = "") -> bool:
        if self._ignorato(valore):
            return False
        parole = [p.strip(".,;:()[]{}\"'’") for p in re.split(r"[\s\-]+", valore.strip()) if p.strip()]
        if len(parole) < 2:
            return False
        if any(p.lower() in _NON_PERSONA for p in parole):
            return False
        if parole[0] in {"San", "Santa", "Santo"}:
            return False
        # Un nome di persona non inizia con preposizione/articolo (es. "DI Milano", "DELLA Corte")
        if parole[0].lower() in _PREP_ARTICOLI:
            return False
        if re.search(r"(?:tribunale|corte|comune|provincia|regione|agenzia|banca|ministero|"
                     r"foro|camera|consiglio|procura|sezione|ordine|collegio|prefettura|questura)"
                     r"(?:\s+\w+){0,2}\s+(?:di|del|della|dello|dei|degli|delle)\s*$",
                     contesto[-40:].lower()):
            return False
        return True

    def _nomi_fallback(self, testo: str) -> str:
        pat = re.compile(rf"\b(?P<nome>{NOME_PERSONA_2_4})\b")
        def seg(s: str) -> str:
            out, ultimo = [], 0
            for m in pat.finditer(s):
                cand = m.group("nome")
                if self._sembra_nome(cand, s[max(0, m.start()-50):m.start()]):
                    out.append(s[ultimo:m.start()]); out.append(self._segnaposto("PERSONA", cand)); ultimo = m.end()
            out.append(s[ultimo:])
            return "".join(out)
        return self._fuori_placeholder(testo, seg)

    def pseudonimizza(self, testo: str) -> str:
        # 0) associazioni definite dall'utente: tutte le forme -> stesso segnaposto.
        if not self._cat_esclusa("SOGGETTO"):
            testo = self._applica_gruppi(testo)
        # 1) individua gli alias introdotti nel testo (denominato «X», di seguito X...)
        #    e i nomi di società senza sigla (da "Alfa Beta S.r.l." -> "Alfa Beta").
        if not self._cat_esclusa("DENOMINAZIONE"):
            self._cattura_alias(testo)
        if not self._cat_esclusa("SOCIETA"):
            self._cattura_nomi_societa(testo)
        # 2) pattern strutturati (società, C.F., ecc.), NER, fallback nomi
        #    (le categorie escluse vengono saltate dentro _sostituisci).
        for d in PATTERN_STRUTTURATI:
            testo = self._sostituisci(testo, d)
        if not self._cat_esclusa("PERSONA"):
            testo = self._nomi_ner(testo)
            if self.fallback_nomi:
                testo = self._nomi_fallback(testo)
        # 3) termini decisi dall'utente o alias: applicati DOPO, così non spezzano
        #    denominazioni più lunghe già riconosciute (es. "Pippo Costruzioni S.r.l.").
        testo = self._applica_termini_forzati(testo)
        if self.pseudonimizza_date:
            for d in PATTERN_DATE_GENERICHE:
                testo = self._sostituisci(testo, d)
        return testo

    def audit_residui(self, testo: str, max_esempi: int = 10) -> Dict[str, List[str]]:
        residui: Dict[str, List[str]] = {}
        def cerca(s: str, d: PatternDef) -> str:
            es = residui.setdefault(d.categoria, [])
            for m in re.finditer(d.pattern, s, flags=d.flags):
                # FIX v4: controllo PRIMA dell'append (il limite poteva essere superato).
                if len(es) >= max_esempi:
                    break
                v = re.sub(r"\s+", " ", m.group(0)).strip()
                if v and v not in es:
                    es.append(v[:180])
            return s
        for d in AUDIT_PATTERNS:
            # Non segnalare come "residuo" ciò che l'utente ha scelto di NON pseudonimizzare.
            if any(d.categoria.startswith(cat) for cat in self.categorie_escluse):
                continue
            self._fuori_placeholder(testo, lambda s, d=d: cerca(s, d))
        if not self._cat_esclusa("PERSONA"):
            nomi: List[str] = []
            pat = re.compile(rf"\b(?P<nome>{NOME_PERSONA_2_4})\b")
            def an(s: str) -> str:
                for m in pat.finditer(s):
                    if len(nomi) >= max_esempi:
                        break
                    c = m.group("nome")
                    if self._sembra_nome(c, s[max(0, m.start()-50):m.start()]) and c not in nomi:
                        nomi.append(c[:180])
                return s
            self._fuori_placeholder(testo, an)
            if nomi: residui["POSSIBILE_PERSONA_RESIDUA"] = nomi
        return {k: v for k, v in residui.items() if v}

    def candidati_dubbi(self, testo: str, max_cand: int = 40) -> List[str]:
        """Termini incerti da sottoporre all'utente: nomi possibili residui e
        parole con iniziale maiuscola che si ripetono (possibili denominazioni)."""
        residui = self.audit_residui(testo)
        cand = list(residui.get("POSSIBILE_PERSONA_RESIDUA", []))
        conta: Dict[str, int] = {}
        def seg(s: str) -> str:
            for m in re.finditer(rf"\b[{_MAIU}][{_LETT}]{{2,}}\b", s):
                w = m.group(0)
                if w.lower() in _NON_PERSONA or w.lower() in _PREP_ARTICOLI:
                    continue
                conta[w] = conta.get(w, 0) + 1
            return s
        self._fuori_placeholder(testo, seg)
        for w, c in conta.items():
            if c >= 2:
                cand.append(w)
        out, visti = [], set()
        for c in cand:
            n = _norm_termine(c)
            if n in visti or n in self.termini_forzati or n in self.termini_ignora:
                continue
            visti.add(n); out.append(c)
            if len(out) >= max_cand:
                break
        return out


# ---- revisione interattiva dei termini dubbi ----
def esegui_revisione(anon: "Pseudonimizzatore", testo: str, chiedi_dubbio, decisioni) -> str:
    """Chiede all'utente, per ogni termine dubbio, se pseudonimizzarlo.
    'chiedi_dubbio(term)' deve restituire 'si', 'no' oppure 'stop'.
    Le scelte vengono ricordate (decisioni globali) e riapplicate."""
    cand = anon.candidati_dubbi(testo)
    if not cand:
        return testo
    modificato = False
    for term in cand:
        try:
            risposta = chiedi_dubbio(term)
        except Exception:
            break
        if risposta == "stop":
            break
        if risposta == "si":
            anon.aggiungi_termine_forzato(term, "DATO_SENSIBILE")
            decisioni["pseudonimizza"][_norm_termine(term)] = {"categoria": "DATO_SENSIBILE", "testo": term}
            modificato = True
        elif risposta == "no":
            anon.aggiungi_ignora(term)
            n = _norm_termine(term)
            if n not in decisioni["ignora"]:
                decisioni["ignora"].append(n)
    if modificato:
        testo = anon._applica_termini_forzati(testo)
    salva_decisioni(decisioni)
    return testo

def _chiedi_dubbio_console(term: str) -> str:
    while True:
        r = input(f"  «{term}»  è un dato sensibile? [s]ì / [n]o / [t]ermina: ").strip().lower()
        if r in ("s", "si", "sì", "y"): return "si"
        if r in ("n", "no"):            return "no"
        if r in ("t", "stop", "q"):     return "stop"
        print("     Rispondi s, n oppure t.")

def analizza_associazioni(righe) -> List[dict]:
    """Trasforma righe di testo (una per soggetto, forme separate da virgola)
    in gruppi di associazioni."""
    gruppi = []
    for riga in righe:
        forme = [f.strip() for f in str(riga).split(",") if f.strip()]
        if forme:
            gruppi.append({"categoria": "SOGGETTO", "forme": forme})
    return gruppi

def scegli_categorie_escluse_console() -> set:
    """Mostra le categorie e consente di indicare quelle da NON pseudonimizzare."""
    print("\n--- Categorie da NON pseudonimizzare ---")
    for i, (c, et) in enumerate(CATEGORIE_PSEUDONIMIZZABILI, 1):
        print(f"  {i:>2}) {et}  [{c}]")
    print("Digita i NUMERI (o i codici) separati da virgola. INVIO per nessuna.")
    x = input("Escludi: ").strip()
    if not x:
        return set()
    escluse = set()
    for tok in re.split(r"[,\s]+", x):
        tok = tok.strip()
        if not tok:
            continue
        if tok.isdigit():
            i = int(tok)
            if 1 <= i <= len(CATEGORIE_PSEUDONIMIZZABILI):
                escluse.add(CATEGORIE_PSEUDONIMIZZABILI[i - 1][0])
        elif tok.upper() in _CODICI_CATEGORIA:
            escluse.add(tok.upper())
    if escluse:
        print("  Non verranno pseudonimizzate: " +
              ", ".join(_ETICHETTA_CATEGORIA[c] for c in sorted(escluse)))
    return escluse

def modifica_associazioni_console(righe_esistenti: List[str]) -> List[str]:
    """Mostra le associazioni già salvate e consente di rimuoverle/aggiungerne.
    Restituisce l'elenco aggiornato di righe (una per soggetto)."""
    righe = list(righe_esistenti)
    print("\n--- Associazioni di nomi (le forme di uno stesso soggetto) ---")
    print("Ogni riga è un soggetto; le forme sono separate da virgola e diventano")
    print("lo stesso segnaposto. Esempio:")
    print("  Alfa Beta Costruzioni S.r.l., Alfa Beta, Alfa, la Società")
    while True:
        print("\nAssociazioni attuali:")
        if righe:
            for i, r in enumerate(righe, 1):
                print(f"  {i}) {r}")
        else:
            print("  (nessuna)")
        print("Comandi: [numero] rimuovi · [testo con virgole] aggiungi · INVIO per confermare")
        x = input("> ").strip()
        if not x:
            break
        if x.isdigit():
            i = int(x)
            if 1 <= i <= len(righe):
                rimossa = righe.pop(i - 1)
                print(f"  rimossa: {rimossa}")
            else:
                print("  numero non valido.")
        else:
            forme = [f for f in x.split(",") if f.strip()]
            if len(forme) >= 1:
                righe.append(x)
                print(f"  aggiunta ({len(forme)} forme).")
    return righe


# ---- I/O ----
def leggi_txt(p: Path) -> str:
    # FIX v4: cp1252 PRIMA di latin-1 (latin-1 non fallisce mai: il ramo cp1252
    # era irraggiungibile e ’ – ecc. venivano decodificati come caratteri di
    # controllo, rompendo i pattern sui nomi con apostrofo tipografico).
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return p.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Impossibile decodificare il file di testo.")

def _ocr_pagine_pdf(percorso: Path, testi: List[str], indici: List[int]) -> List[str]:
    """Rasterizza e riconosce via OCR le pagine PDF prive di testo estraibile."""
    import io
    try:
        doc = fitz.open(str(percorso))
    except Exception:
        return testi
    lang = _lingua_ocr()
    kw = {"lang": lang} if lang else {}
    for i in indici:
        try:
            pix = doc.load_page(i).get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            t = pytesseract.image_to_string(img, **kw)
            if t and t.strip():
                testi[i] = t
                print(f"[OCR] Pagina {i + 1}: testo riconosciuto tramite OCR.")
        except Exception:
            pass
    try:
        doc.close()
    except Exception:
        pass
    return testi

def leggi_pdf(p: Path, force_ocr: bool = False) -> str:
    if pypdf is None:
        raise RuntimeError("pypdf non installato. Usa: pip install pypdf")
    reader = pypdf.PdfReader(str(p))
    testi: List[str] = []
    for pagina in reader.pages:
        if force_ocr:
            testi.append("")
            continue
        try:
            t = pagina.extract_text(extraction_mode="layout")
        except Exception:
            t = pagina.extract_text() or ""
        testi.append(t or "")
    # Pagine senza testo: probabile scansione -> OCR (se disponibile).
    da_ocr = [i for i, t in enumerate(testi) if not t.strip()]
    if da_ocr and ocr_disponibile_pdf():
        testi = _ocr_pagine_pdf(p, testi, da_ocr)
    parti = [f"\n--- PAGINA {i + 1} ---\n{t}" for i, t in enumerate(testi)]
    testo = "\n".join(parti)
    if not testo.strip():
        if not ocr_disponibile_pdf():
            raise ValueError(
                "PDF senza testo (probabile scansione) e OCR non disponibile. "
                "Installa: pip install pytesseract pymupdf pillow, più il motore "
                "Tesseract con la lingua italiana. Vedi il manuale.")
        raise ValueError("PDF vuoto: nessun testo riconosciuto nemmeno con l'OCR.")
    return testo

def leggi_immagine(p: Path) -> str:
    if not ocr_disponibile_immagini():
        raise ValueError(
            "Per leggere le immagini serve l'OCR. Installa: pip install pytesseract "
            "pillow, più il motore Tesseract con la lingua italiana. Vedi il manuale.")
    try:
        img = Image.open(str(p))
    except Exception as exc:
        raise ValueError(f"Impossibile aprire l'immagine: {exc}") from exc
    testo = _ocr_immagine(img)
    if not testo.strip():
        raise ValueError("Nessun testo riconosciuto nell'immagine.")
    return testo

def leggi_docx(p: Path) -> str:
    if _python_docx is None:
        raise ValueError("Per i file Word (.docx) serve: pip install python-docx")
    try:
        d = _python_docx.Document(str(p))
    except Exception as exc:
        raise ValueError(f"Impossibile aprire il documento Word: {exc}") from exc
    parti: List[str] = []
    for para in d.paragraphs:
        if para.text:
            parti.append(para.text)
    # Anche il testo nelle tabelle (frequenti negli atti).
    for tab in d.tables:
        for row in tab.rows:
            celle = [c.text.strip() for c in row.cells]
            if any(celle):
                parti.append("\t".join(celle))
    testo = "\n".join(parti)
    if not testo.strip():
        raise ValueError("Documento Word vuoto o senza testo estraibile.")
    return testo

def carica_contenuto(p: Path, force_ocr: bool = False) -> str:
    ext = p.suffix.lower()
    if ext == ".pdf": return leggi_pdf(p, force_ocr=force_ocr)
    if ext == ".txt": return leggi_txt(p)
    if ext == ".docx": return leggi_docx(p)
    if ext in _ESTENSIONI_IMG: return leggi_immagine(p)
    if ext == ".doc":
        raise ValueError("Formato .doc (vecchio Word) non supportato: salvalo come .docx e riprova.")
    raise ValueError(f"Formato non supportato: {ext}")

_ESTENSIONI_TESTO = {".txt", ".pdf", ".docx"}

def file_elaborabile(p: Path) -> bool:
    if p.suffix.lower() not in (_ESTENSIONI_TESTO | _ESTENSIONI_IMG): return False
    n = p.name.lower()
    return not any(n.endswith(e) for e in (
        "_pseudonimizzato.txt","_ripristinato.txt","_residui_da_verificare.txt","_associazioni.txt",
        "_mappa_privata.json","_mappa_privata.json.enc",
        "_report_privacy.json","mappa_fascicolo_privata.json","mappa_fascicolo_privata.json.enc",
        "report_fascicolo_privacy.json"))

class ErroreUso(Exception):
    """Errore d'uso da parte dell'utente: mostra un messaggio chiaro + l'help."""


def risolvi_bersagli(arg: Optional[str]) -> Tuple[List[Path], Optional[Path]]:
    if arg:
        p = Path(arg).expanduser().resolve()
        if p.is_dir():
            files = sorted([f for f in p.iterdir() if f.is_file() and file_elaborabile(f)])
            if not files:
                raise ErroreUso(f"La cartella '{p}' non contiene file .txt o .pdf da elaborare.")
            return files, p
        if p.is_file():
            if p.suffix.lower() not in (_ESTENSIONI_TESTO | _ESTENSIONI_IMG):
                raise ErroreUso(f"Formato non supportato: '{p.suffix or p.name}'. "
                                "Sono ammessi file .txt, .pdf, .docx o immagini (png, jpg, tiff...).")
            return [p], None
        raise ErroreUso(f"File o cartella inesistente: '{p}'. "
                        "Controlla il nome e il percorso.")
    for nome in ("atto_originale.txt", "atto_originale.pdf"):
        p = Path(nome).resolve()
        if p.exists(): return [p], None
    raise ErroreUso("Nessun file indicato e nessun 'atto_originale.txt/.pdf' trovato "
                    "nella cartella corrente. Indica il file da elaborare.")


_ETICHETTE_RESIDUI = {
    "CODICE_FISCALE_RESIDUO": "Codici fiscali",
    "IBAN_RESIDUO": "IBAN",
    "EMAIL_RESIDUA": "Email",
    "PARTITA_IVA_RESIDUA": "Partite IVA",
    "TELEFONO_RESIDUO": "Telefoni",
    "TARGA_RESIDUA": "Targhe",
    "INDIRIZZO_RESIDUO": "Indirizzi",
    "DATA_NASCITA_RESIDUA": "Date di nascita",
    "LUOGO_NASCITA_RESIDUO": "Luoghi di nascita",
    "DATI_CATASTALI_RESIDUI": "Dati catastali",
    "POSSIBILE_PERSONA_RESIDUA": "Possibili nomi di persona",
}

def formatta_residui(nome_atto: str, residui: Dict[str, List[str]]) -> str:
    tot = sum(len(v) for v in residui.values())
    righe = ["=" * 70,
             "RESIDUI DA VERIFICARE MANUALMENTE",
             f"Atto: {nome_atto}",
             f"Generato: {datetime.now().isoformat(timespec='seconds')}",
             "=" * 70, ""]
    if not tot:
        righe.append("Nessun residuo evidente rilevato dall'audit automatico.")
        righe.append("")
        righe.append("NOTA: l'audit non garantisce l'assenza totale di dati sensibili.")
        righe.append("Consigliata comunque una rilettura del file _pseudonimizzato.txt.")
        return "\n".join(righe) + "\n"
    righe.append(f"Trovati {tot} possibili residui, suddivisi per categoria.")
    righe.append("Controlla nel file _pseudonimizzato.txt se vanno rimossi a mano.")
    righe.append("")
    for cat, esempi in residui.items():
        titolo = _ETICHETTE_RESIDUI.get(cat, cat)
        righe.append(f"[{titolo}] ({len(esempi)})")
        for e in esempi:
            righe.append(f"  - {e}")
        righe.append("")
    return "\n".join(righe) + "\n"


# Categorie che l'utente può scegliere di NON pseudonimizzare (codice + etichetta).
CATEGORIE_PSEUDONIMIZZABILI = [
    ("PERSONA", "Nomi di persona"),
    ("SOCIETA", "Società / denominazioni"),
    ("DENOMINAZIONE", "Alias/denominazioni brevi"),
    ("SOGGETTO", "Associazioni di nomi"),
    ("CONDOMINIO", "Condominio"),
    ("CODICE_FISCALE", "Codice fiscale"),
    ("PARTITA_IVA", "Partita IVA"),
    ("IBAN", "IBAN"),
    ("CONTO_CORRENTE", "Conto corrente"),
    ("CARTA_CREDITO", "Carta di credito"),
    ("IMPORTO", "Importi in euro"),
    ("REPERTORIO_NOTARILE", "Repertorio notarile"),
    ("EMAIL", "Email"),
    ("PEC", "PEC"),
    ("TELEFONO", "Telefono / cellulare / fax"),
    ("DOCUMENTO_IDENTITA", "Documento d'identità"),
    ("TESSERA_SANITARIA", "Tessera sanitaria"),
    ("TARGA", "Targa veicolo"),
    ("DATA_NASCITA", "Data di nascita"),
    ("LUOGO_NASCITA", "Luogo di nascita"),
    ("INDIRIZZO", "Indirizzo"),
    ("CAP", "CAP (codice postale)"),
    ("COORDINATA", "Coordinate geografiche"),
    ("NUMERO_RUOLO", "Numero di ruolo (R.G.)"),
    ("NUMERO_PROVVEDIMENTO", "Numero di provvedimento"),
    ("DATI_CATASTALI", "Dati catastali"),
    ("DATA", "Date generiche (con --date)"),
]
_CODICI_CATEGORIA = {c for c, _ in CATEGORIE_PSEUDONIMIZZABILI}
_ETICHETTA_CATEGORIA = dict(CATEGORIE_PSEUDONIMIZZABILI)

def parse_categorie(s) -> Tuple[set, List[str]]:
    """Da 'DATI_CATASTALI, targa' -> ({'DATI_CATASTALI','TARGA'}, [ignoti])."""
    escluse, ignoti = set(), []
    for tok in re.split(r"[,\s]+", (s or "").strip()):
        if not tok:
            continue
        t = tok.strip().upper()
        if t in _CODICI_CATEGORIA:
            escluse.add(t)
        else:
            ignoti.append(tok)
    return escluse, ignoti


def elabora(bersagli, cartella, pseudonimizza_date, fallback_nomi, encrypt_map,
            force_ocr=False, ottieni_password=None, revisione=False, chiedi_dubbio=None,
            associazioni_righe=None, categorie_escluse=None, catasto_fittizio=False):
    if not bersagli:
        print("[ERRORE] Nessun file da elaborare."); return
    # FIX v4: dipendenza e password verificate PRIMA di elaborare qualsiasi file.
    # ottieni_password: callback per ambienti senza console (es. GUI); se assente
    # si usa la richiesta da terminale.
    password: Optional[str] = None
    if encrypt_map:
        _crypto()
        password = ottieni_password() if ottieni_password is not None else chiedi_password_cifratura()
        if not password:
            print("[ERRORE] Cifratura annullata: nessuna password fornita."); return
    anon = Pseudonimizzatore(pseudonimizza_date, fallback_nomi)
    anon.categorie_escluse = set(categorie_escluse or [])
    if anon.categorie_escluse:
        etichette = ", ".join(_ETICHETTA_CATEGORIA.get(c, c) for c in sorted(anon.categorie_escluse))
        print(f"[NOTA] Categorie NON pseudonimizzate (escluse su tua richiesta): {etichette}.")
    if catasto_fittizio and "DATI_CATASTALI" not in anon.categorie_escluse:
        anon.categorie_fittizie.add("DATI_CATASTALI")
        print("[NOTA] Dati catastali sostituiti con valori FITTIZI coerenti (reversibili con la mappa).")
    # Decisioni globali dell'utente (memoria persistente): applicate SEMPRE.
    decisioni = carica_decisioni()
    for n, info in decisioni["pseudonimizza"].items():
        anon.termini_forzati[n] = (info.get("categoria", "DATO_SENSIBILE"), info.get("testo", n))
    anon.termini_ignora.update(decisioni["ignora"])
    # Associazioni: file affiancato all'atto (stesso nome). Se l'utente ne ha
    # fornite di nuove/modificate si salvano; altrimenti si riusano quelle salvate.
    assoc_file = percorso_associazioni(bersagli, cartella)
    if associazioni_righe is not None:
        nome_atto = cartella.name if cartella else bersagli[0].name
        scrivi_associazioni_txt(assoc_file, associazioni_righe, nome_atto)
        righe_assoc = list(associazioni_righe)
    else:
        righe_assoc = leggi_associazioni_txt(assoc_file)
    anon.gruppi = analizza_associazioni(righe_assoc)
    if anon.gruppi:
        print(f"[NOTA] {len(anon.gruppi)} associazioni di nomi attive (file: {assoc_file.name}).")
    report = {"created_at": datetime.now().isoformat(timespec="seconds"),
              "spacy_attivo": _NLP is not None, "fallback_nomi": fallback_nomi,
              "date_generiche": pseudonimizza_date, "ocr": ocr_disponibile_pdf(),
              "force_ocr": force_ocr, "files": {}}
    errori: Dict[str, str] = {}
    print("=" * 70); print("PSEUDONIMIZZATORE ATTI LEGALI v7"); print("=" * 70)
    stampa_stato_dipendenze()
    for p in bersagli:
        print(f"\n[INFO] Elaboro: {p.name}")
        # FIX v4: un file che fallisce (PDF scansionato, pypdf mancante...) non
        # abortisce più l'intero fascicolo prima del salvataggio della mappa.
        try:
            testo = anon.pseudonimizza(carica_contenuto(p, force_ocr=force_ocr))
            if revisione and chiedi_dubbio is not None:
                testo = esegui_revisione(anon, testo, chiedi_dubbio, decisioni)
            residui = anon.audit_residui(testo)
        except Exception as exc:
            errori[p.name] = str(exc)
            report["files"][p.name] = {"errore": str(exc)}
            print(f"[ERRORE] {p.name}: {exc} — file saltato.")
            continue
        out = p.with_name(p.stem + "_pseudonimizzato.txt"); out.write_text(testo, encoding="utf-8")
        res_file = p.with_name(p.stem + "_residui_da_verificare.txt")
        # FIX v4: il file dei residui contiene dati reali -> permessi 0600.
        scrivi_file_privato(res_file, formatta_residui(p.name, residui))
        report["files"][p.name] = {"output": out.name, "residui_file": res_file.name,
                                   "residui_potenziali": residui}
        print(f"[OK] {out.name}")
        tot = sum(len(v) for v in residui.values())
        if tot:
            print(f"[ATTENZIONE] {tot} possibili residui da verificare -> {res_file.name}")
        else:
            print(f"[OK] Audit: nessun residuo evidente (dettagli in {res_file.name}).")
    if cartella:
        base = cartella / f"{cartella.name}_mappa_privata.json"
        rep = cartella / f"{cartella.name}_report_privacy.json"
    else:
        base = bersagli[0].with_name(bersagli[0].stem + "_mappa_privata.json")
        rep = bersagli[0].with_name(bersagli[0].stem + "_report_privacy.json")
    mp = salva_mappa(anon.mappa, base, password)
    report["mappa"] = {"file": mp.name, "cifrata": password is not None,
                       "segnaposto": len(anon.mappa), "per_categoria": anon.contatori}
    if errori:
        report["errori"] = errori
    scrivi_file_privato(rep, json.dumps(report, ensure_ascii=False, indent=2))
    print("\n" + "-" * 70)
    print(f"[OK] Mappa privata: {mp.name} ({len(anon.mappa)} sostituzioni)")
    print(f"[OK] Report: {rep.name}")
    if errori:
        print(f"[ATTENZIONE] {len(errori)} file saltati per errori: {', '.join(errori)}")
    print("[AVVISO] Mappa, report e file _residui contengono dati reali:")
    print("         non inviarli all'AI, conservali offline.")
    print("=" * 70)
    stampa_disclaimer()
    stampa_crediti()


def ripristina(risposta_path, mappa_path, output=None):
    risposta = Path(risposta_path).expanduser().resolve()
    mappa = carica_mappa(Path(mappa_path).expanduser().resolve())
    testo = risposta.read_text(encoding="utf-8")
    # 1) segnaposto tra parentesi quadre ([PERSONA_1], ...)
    testo = PLACEHOLDER_RE.sub(lambda m: mappa.get(m.group(0), m.group(0)), testo)
    # 2) dati fittizi coerenti (chiavi NON tra parentesi, es. "Foglio 51"):
    #    ripristina il valore reale, dal più lungo al più corto.
    fittizi = [(k, v) for k, v in mappa.items() if not PLACEHOLDER_RE.fullmatch(k)]
    for k, v in sorted(fittizi, key=lambda kv: len(kv[0]), reverse=True):
        testo = re.sub(rf"(?<![\wÀ-ÿ]){re.escape(k)}(?![\wÀ-ÿ])",
                       lambda _m, v=v: v, testo)
    # FIX v4: conserva l'estensione originale della risposta (es. .md).
    out = (Path(output).expanduser().resolve() if output
           else risposta.with_name(risposta.stem + "_ripristinato" + (risposta.suffix or ".txt")))
    out.write_text(testo, encoding="utf-8")
    print(f"[OK] Risposta ripristinata: {out}")
    stampa_crediti()


class _FormatterIT(argparse.HelpFormatter):
    def add_usage(self, usage, actions, groups, prefix=None):
        return super().add_usage(usage, actions, groups, prefix or "uso: ")

def _input_nonvuoto(prompt: str) -> str:
    while True:
        v = input(prompt).strip()
        if v:
            return v
        print("  Inserisci un valore (o Ctrl+C per uscire).")

def _input_file_esistente(prompt: str) -> str:
    while True:
        v = _input_nonvuoto(prompt)
        if Path(v).expanduser().exists():
            return v
        print(f"  [!] Percorso inesistente: {v}")

def _chiedi_si_no(prompt: str, default: bool = False) -> bool:
    suff = " [S/n] " if default else " [s/N] "
    while True:
        v = input(prompt + suff).strip().lower()
        if not v:
            return default
        if v in ("s", "si", "sì", "y", "yes"):
            return True
        if v in ("n", "no"):
            return False
        print("  Rispondi 's' oppure 'n'.")

def _processi_windows() -> Dict[int, Tuple[int, str]]:
    """Mappa pid -> (pid_padre, nome_exe) di tutti i processi (solo Windows)."""
    import ctypes
    from ctypes import wintypes
    class PROCESSENTRY32(ctypes.Structure):
        _fields_ = [("dwSize", wintypes.DWORD), ("cntUsage", wintypes.DWORD),
                    ("th32ProcessID", wintypes.DWORD),
                    ("th32DefaultHeapID", ctypes.POINTER(ctypes.c_ulong)),
                    ("th32ModuleID", wintypes.DWORD), ("cntThreads", wintypes.DWORD),
                    ("th32ParentProcessID", wintypes.DWORD),
                    ("pcPriClassBase", ctypes.c_long), ("dwFlags", wintypes.DWORD),
                    ("szExeFile", ctypes.c_char * 260)]
    k = ctypes.windll.kernel32
    snap = k.CreateToolhelp32Snapshot(0x00000002, 0)  # TH32CS_SNAPPROCESS
    entry = PROCESSENTRY32(); entry.dwSize = ctypes.sizeof(PROCESSENTRY32)
    info: Dict[int, Tuple[int, str]] = {}
    if k.Process32First(snap, ctypes.byref(entry)):
        while True:
            info[entry.th32ProcessID] = (entry.th32ParentProcessID,
                                         entry.szExeFile.decode(errors="ignore").lower())
            if not k.Process32Next(snap, ctypes.byref(entry)):
                break
    k.CloseHandle(snap)
    return info

def _avviato_da_shell_windows() -> bool:
    """True se il programma è stato lanciato da un terminale (cmd, PowerShell...)."""
    _SHELL = {"cmd.exe", "powershell.exe", "pwsh.exe", "wt.exe",
              "windowsterminal.exe", "bash.exe", "code.exe"}
    try:
        info = _processi_windows()
        mio = os.getpid()
        mio_nome = info.get(mio, (0, ""))[1]
        pid = info.get(mio, (0, ""))[0]
        visti = set()
        while pid and pid not in visti:
            visti.add(pid)
            ppid, nome = info.get(pid, (0, ""))
            if nome and nome != mio_nome:        # salta il bootloader PyInstaller
                if nome in _SHELL:
                    return True                  # avviato da una shell
                return False                     # explorer.exe o altro: doppio clic
            pid = ppid
    except Exception:
        pass
    return False

def _serve_pausa() -> bool:
    """Decide se attendere INVIO: solo quando il programma ha una finestra propria
    (doppio clic), non quando è lanciato da un terminale già aperto."""
    if os.environ.get("ANON_ATTI_NO_PAUSE"):
        return False
    if os.environ.get("ANON_ATTI_PAUSE"):
        return True
    try:
        if not sys.stdin.isatty() or not sys.stdout.isatty():
            return False                         # output rediretto/pipe: niente pausa
    except Exception:
        return False
    if os.name == "nt":
        return not _avviato_da_shell_windows()
    # macOS/Linux: non è possibile distinguere in modo affidabile il doppio clic
    # dal Terminale; il file .command apre comunque una finestra propria, quindi
    # la pausa è utile. Chi lancia da shell può disattivarla con ANON_ATTI_NO_PAUSE=1.
    return True

def _pausa_finale() -> None:
    """Tiene aperta la finestra quando il programma è avviato con un clic."""
    if not _serve_pausa():
        return
    try:
        input("\nPremi INVIO per chiudere...")
    except (EOFError, KeyboardInterrupt):
        pass

def modalita_interattiva() -> int:
    print("=" * 70)
    print("PSEUDONIMIZZATORE ATTI LEGALI v7 — modalità interattiva")
    print("=" * 70)
    print("Cosa vuoi fare?")
    print("  1) Pseudonimizzare un atto (o una cartella di atti)")
    print("  2) Ripristinare una risposta già pseudonimizzata")

    while True:
        scelta = input("Scelta [1/2]: ").strip()
        if scelta in ("1", "2"):
            break
        print("  Digita 1 o 2.")

    if scelta == "1":
        # Chiede il file/cartella e riprova in caso di percorso errato,
        # senza chiudere il programma.
        print("\nPuoi indicare DUE cose diverse:")
        print("  1. UN SINGOLO FILE  -> viene pseudonimizzato solo quel documento.")
        print("     Esempi:")
        print("       Windows:      C:\\Users\\Mario\\Documenti\\atto.pdf")
        print("       macOS/Linux:  /Users/mario/Documenti/atto.pdf")
        print("  2. UNA CARTELLA     -> vengono pseudonimizzati TUTTI i file")
        print("     (.txt, .pdf, immagini) contenuti al suo interno, con un'unica")
        print("     mappa per l'intero fascicolo.")
        print("     Esempi:")
        print("       Windows:      C:\\Users\\Mario\\Documenti\\Fascicolo_Rossi")
        print("       macOS/Linux:  /Users/mario/Documenti/Fascicolo_Rossi")
        print("Suggerimento: puoi anche TRASCINARE il file o la cartella in questa")
        print("finestra e premere INVIO: il percorso viene inserito da solo.")
        while True:
            percorso = _input_nonvuoto("\nFile o cartella da pseudonimizzare: ")
            try:
                bersagli, cartella = risolvi_bersagli(percorso)
                break
            except ErroreUso as exc:
                print(f"  [!] {exc}")

        if cartella:
            print(f"\n[INFO] Cartella riconosciuta: {len(bersagli)} file da elaborare.")
        else:
            print("\n[INFO] File singolo pronto per l'elaborazione.")
        pseudonimizza_date = _chiedi_si_no("Pseudonimizzare anche le date generiche?", default=False)
        fallback_nomi   = _chiedi_si_no("Usare il fallback regex per i nomi 'nudi'?", default=True)
        encrypt_map     = _chiedi_si_no("Cifrare la mappa con password?", default=False)
        force_ocr = False
        if ocr_disponibile_pdf():
            force_ocr = _chiedi_si_no("Forzare l'OCR su tutte le pagine dei PDF?", default=False)
        revisione = _chiedi_si_no("Rivedere a mano i termini dubbi (e ricordare le scelte)?", default=False)
        associazioni_righe = None
        esistenti = leggi_associazioni_txt(percorso_associazioni(bersagli, cartella))
        domanda = ("Modificare le associazioni di nomi già salvate?" if esistenti
                   else "Indicare associazioni di nomi (più forme dello stesso soggetto)?")
        if _chiedi_si_no(domanda, default=bool(esistenti)):
            associazioni_righe = modifica_associazioni_console(esistenti)
        categorie_escluse = set()
        if _chiedi_si_no("Vuoi escludere qualche categoria dalla pseudonimizzazione?", default=False):
            categorie_escluse = scegli_categorie_escluse_console()
        catasto_fittizio = False
        if "DATI_CATASTALI" not in categorie_escluse:
            catasto_fittizio = _chiedi_si_no(
                "Sostituire i dati catastali con valori FITTIZI coerenti (invece del segnaposto)?", default=False)

        elabora(bersagli, cartella, pseudonimizza_date, fallback_nomi, encrypt_map, force_ocr,
                revisione=revisione,
                chiedi_dubbio=_chiedi_dubbio_console if revisione else None,
                associazioni_righe=associazioni_righe,
                categorie_escluse=categorie_escluse,
                catasto_fittizio=catasto_fittizio)
        return 0

    # scelta == "2": ripristino
    print("\n--- Ripristino ---")
    risposta = _input_file_esistente("File con la risposta dell'AI (contiene i segnaposto): ")
    mappa    = _input_file_esistente("File mappa (_mappa_privata.json oppure .json.enc): ")
    out      = input("File di output (INVIO per il nome automatico): ").strip() or None

    ripristina(risposta, mappa, out)
    return 0

# ---- interfaccia grafica (tkinter, libreria standard) ----
def _gui_disponibile() -> bool:
    """True se è possibile aprire una finestra grafica su questo sistema."""
    try:
        import tkinter  # noqa: F401
    except Exception:
        return False
    # Su Linux serve un server grafico attivo (variabile DISPLAY/WAYLAND).
    if sys.platform.startswith("linux") and not (
            os.environ.get("DISPLAY") or os.environ.get("WAYLAND_DISPLAY")):
        return False
    return True

def avvia_gui() -> int:
    """Finestra grafica per chi fatica a trovare il percorso dei file.

    Usa solo tkinter, presente su Windows e installabile su Linux con
    'sudo apt install python3-tk'. I pulsanti "Sfoglia" aprono il file-browser
    nativo del sistema operativo, così l'utente naviga con il mouse senza mai
    digitare un percorso. La finestra riusa la stessa pipeline della riga di
    comando (risolvi_bersagli / elabora / ripristina): il comportamento e i
    file prodotti sono identici.
    """
    try:
        import tkinter as tk
        from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
    except Exception:
        print("[ERRORE] Interfaccia grafica non disponibile: manca il modulo tkinter.\n"
              "         Su Linux installa:  sudo apt install python3-tk\n"
              "         In alternativa usa la riga di comando (vedi --aiuto).",
              file=sys.stderr)
        return 3

    import queue, threading, platform, subprocess
    from contextlib import redirect_stdout, redirect_stderr

    radice = tk.Tk()
    radice.title("Pseudonimizzatore Atti Legali")
    radice.geometry("760x660")
    radice.minsize(680, 560)

    # Tooltip discreti (compaiono dopo qualche secondo sul widget).
    def aggiungi_tooltip(widget, testo, ritardo=1200):
        stato = {"id": None, "tip": None}
        def mostra():
            if stato["tip"] or not widget.winfo_exists():
                return
            x = widget.winfo_rootx() + 18
            y = widget.winfo_rooty() + widget.winfo_height() + 6
            tip = tk.Toplevel(widget)
            tip.wm_overrideredirect(True)
            tip.wm_geometry(f"+{x}+{y}")
            try:
                tip.attributes("-topmost", True)
            except Exception:
                pass
            tk.Label(tip, text=testo, bg="#ffffe0", fg="#333333", relief="solid", bd=1,
                     justify="left", wraplength=320, font=("TkDefaultFont", 9)).pack(ipadx=5, ipady=3)
            stato["tip"] = tip
        def pianifica(_=None):
            annulla()
            stato["id"] = widget.after(ritardo, mostra)
        def annulla(_=None):
            if stato["id"]:
                widget.after_cancel(stato["id"]); stato["id"] = None
            if stato["tip"]:
                stato["tip"].destroy(); stato["tip"] = None
        widget.bind("<Enter>", pianifica, add="+")
        widget.bind("<Leave>", annulla, add="+")
        widget.bind("<ButtonPress>", annulla, add="+")

    percorso_var    = tk.StringVar()
    date_var        = tk.BooleanVar(value=False)
    cifra_var       = tk.BooleanVar(value=False)   # password non richiesta di default
    no_fallback_var = tk.BooleanVar(value=False)
    revisiona_var   = tk.BooleanVar(value=False)
    catasto_fitt_var = tk.BooleanVar(value=False)
    risposta_var    = tk.StringVar()
    mappa_var       = tk.StringVar()

    coda_log: "queue.Queue[str]" = queue.Queue()
    password_pre = [None]   # riempita sul thread principale prima di elaborare
    assoc_righe: List[str] = []   # righe di associazioni inserite dall'utente
    assoc_aperto = [False]        # True se l'utente ha aperto/confermato la finestra
    cat_escluse: set = set()      # categorie da NON pseudonimizzare

    # --- log a video ---
    class _ScrittoreCoda:
        def write(self, s):
            for riga in str(s).splitlines():
                if riga.strip():
                    coda_log.put(riga)
        def flush(self):
            pass

    def log(msg):
        coda_log.put(msg)

    def drena_coda():
        try:
            while True:
                msg = coda_log.get_nowait()
                area_log.configure(state="normal")
                area_log.insert("end", msg + "\n")
                area_log.see("end")
                area_log.configure(state="disabled")
        except queue.Empty:
            pass
        radice.after(120, drena_coda)

    # --- selettori nativi ---
    def scegli_file():
        p = filedialog.askopenfilename(
            title="Scegli l'atto da pseudonimizzare",
            filetypes=[("Atti (PDF, Word, TXT)", "*.pdf *.docx *.txt"), ("PDF", "*.pdf"),
                       ("Word", "*.docx"), ("Testo", "*.txt"),
                       ("Immagini", "*.png *.jpg *.jpeg *.tif *.tiff"), ("Tutti i file", "*.*")])
        if p:
            percorso_var.set(p)

    def scegli_cartella():
        p = filedialog.askdirectory(title="Scegli la cartella con gli atti")
        if p:
            percorso_var.set(p)

    def scegli_risposta():
        p = filedialog.askopenfilename(
            title="Scegli la risposta dell'AI da ripristinare",
            filetypes=[("Testo/Markdown", "*.txt *.md"), ("Tutti i file", "*.*")])
        if p:
            risposta_var.set(p)

    def scegli_mappa():
        p = filedialog.askopenfilename(
            title="Scegli la mappa privata",
            filetypes=[("Mappa", "*.json *.enc"), ("Tutti i file", "*.*")])
        if p:
            mappa_var.set(p)

    def apri_cartella():
        perc = percorso_var.get().strip()
        cartella = perc if os.path.isdir(perc) else os.path.dirname(perc)
        if not cartella or not os.path.isdir(cartella):
            log("[INFO] Nessuna cartella valida selezionata.")
            return
        try:
            if platform.system() == "Windows":
                os.startfile(cartella)            # type: ignore[attr-defined]
            elif platform.system() == "Darwin":
                subprocess.Popen(["open", cartella])
            else:
                subprocess.Popen(["xdg-open", cartella])
        except Exception as exc:
            log(f"[ERRORE] Impossibile aprire la cartella: {exc}")

    def _chiedi_password_gui():
        pw1 = simpledialog.askstring("Cifratura mappa",
            "Password per cifrare la mappa:", show="*", parent=radice)
        if pw1 is None:
            return None
        if not pw1.strip():
            messagebox.showerror("Errore", "Password vuota non ammessa.")
            return None
        pw2 = simpledialog.askstring("Cifratura mappa",
            "Conferma la password:", show="*", parent=radice)
        if pw2 is None:
            return None
        if pw1 != pw2:
            messagebox.showerror("Errore", "Le password non coincidono.")
            return None
        return pw1

    def gestisci_associazioni():
        # Pre-carica le associazioni già salvate accanto all'atto selezionato.
        if not assoc_righe:
            ap = assoc_path_da_percorso(percorso_var.get())
            if ap is not None:
                assoc_righe.extend(leggi_associazioni_txt(ap))
        win = tk.Toplevel(radice)
        win.title("Associazioni di nomi")
        win.geometry("580x400")
        win.transient(radice); win.grab_set()
        ttk.Label(win, justify="left", text=(
            "Una riga per ogni soggetto. Elenca le sue forme SEPARATE DA VIRGOLA\n"
            "(nome completo, abbreviazioni, alias): verranno sostituite tutte con\n"
            "lo stesso segnaposto. Vengono salvate accanto all'atto e ri-mostrate\n"
            "alla prossima elaborazione dello stesso file.")
        ).pack(anchor="w", padx=10, pady=(10, 6))
        txt = tk.Text(win, height=12, wrap="word")
        txt.pack(fill="both", expand=True, padx=10)
        if assoc_righe:
            txt.insert("1.0", "\n".join(assoc_righe))
        else:
            # Righe di esempio con '#': restano SOLO come guida e non vengono
            # mai salvate né pseudonimizzate. Scrivi le tue associazioni sotto.
            txt.insert("1.0",
                       "# Esempio (queste righe con # NON vengono usate):\n"
                       "#   Rossi Costruzioni S.r.l., Rossi Costruzioni, la Società\n"
                       "#   Mario Rossi, Rossi, il ricorrente\n"
                       "# Scrivi qui sotto le tue associazioni (una per riga):\n")
            txt.mark_set("insert", "end")
        def ok():
            assoc_righe.clear()
            for r in txt.get("1.0", "end").splitlines():
                r = r.strip()
                if r and not r.startswith("#"):
                    assoc_righe.append(r)
            assoc_aperto[0] = True
            win.destroy()
        barra = ttk.Frame(win); barra.pack(fill="x", padx=10, pady=8)
        ttk.Button(barra, text="OK", command=ok).pack(side="right")
        ttk.Button(barra, text="Annulla", command=win.destroy).pack(side="right", padx=(0, 6))
        win.wait_window()

    def gestisci_categorie():
        win = tk.Toplevel(radice)
        win.title("Categorie da NON pseudonimizzare")
        win.geometry("560x520")
        win.transient(radice); win.grab_set()
        ttk.Label(win, justify="left", text=(
            "Spunta le categorie che vuoi LASCIARE nel testo (non pseudonimizzate).\n"
            "Le altre verranno pseudonimizzate normalmente.")
        ).pack(anchor="w", padx=12, pady=(12, 6))
        cornice = ttk.Frame(win); cornice.pack(fill="both", expand=True, padx=12)
        vars_cat = {}
        for i, (cod, et) in enumerate(CATEGORIE_PSEUDONIMIZZABILI):
            v = tk.BooleanVar(value=(cod in cat_escluse))
            vars_cat[cod] = v
            ttk.Checkbutton(cornice, text=et, variable=v).grid(
                row=i % 11, column=i // 11, sticky="w", padx=6, pady=2)
        def ok():
            cat_escluse.clear()
            for cod, v in vars_cat.items():
                if v.get():
                    cat_escluse.add(cod)
            win.destroy()
        barra = ttk.Frame(win); barra.pack(fill="x", padx=12, pady=10)
        ttk.Button(barra, text="OK", command=ok).pack(side="right")
        ttk.Button(barra, text="Annulla", command=win.destroy).pack(side="right", padx=(0, 6))
        win.wait_window()

    def _chiedi_dubbio_gui(term):
        # Dialogo modale: Sì = pseudonimizza, No = lascia, Annulla = ferma la revisione.
        r = messagebox.askyesnocancel(
            "Revisione dei termini dubbi",
            f"«{term}»\n\nÈ un dato sensibile da pseudonimizzare?\n"
            "(Sì = pseudonimizza sempre · No = lascia sempre · Annulla = interrompi)",
            parent=radice)
        if r is None:
            return "stop"
        return "si" if r else "no"

    # --- esecuzione ---
    # I widget vengono letti SOLO sul thread principale (tkinter non è
    # thread-safe): i valori sono raccolti in on_esegui e passati al worker.
    def _lavora(scheda_ripristino, dati):
        try:
            if scheda_ripristino:                    # scheda "Ripristina"
                r, m = dati["risposta"], dati["mappa"]
                if not r or not m:
                    log("[ERRORE] Seleziona sia la risposta dell'AI sia la mappa.")
                    return
                with redirect_stdout(_ScrittoreCoda()), redirect_stderr(_ScrittoreCoda()):
                    ripristina(r, m)
            else:                                    # scheda "Pseudonimizza"
                bersagli, cartella = risolvi_bersagli(dati["percorso"] or None)
                with redirect_stdout(_ScrittoreCoda()), redirect_stderr(_ScrittoreCoda()):
                    elabora(bersagli, cartella, dati["date"], not dati["no_fallback"],
                            dati["cifra"], ottieni_password=lambda: password_pre[0],
                            revisione=dati["revisiona"],
                            chiedi_dubbio=_chiedi_dubbio_gui if dati["revisiona"] else None,
                            associazioni_righe=dati["associazioni_righe"],
                            categorie_escluse=dati["categorie_escluse"],
                            catasto_fittizio=dati["catasto_fittizio"])
            log("")
            log(">>> Operazione completata. <<<")
        except ErroreUso as exc:
            log(f"[ERRORE] {exc}")
        except Exception as exc:
            log(f"[ERRORE] {exc}")
        finally:
            # after() è sicuro sia dal thread principale sia dal worker.
            radice.after(0, lambda: bottone_esegui.configure(state="normal"))

    def on_esegui():
        bottone_esegui.configure(state="disabled")
        password_pre[0] = None
        scheda = note.index(note.select())
        dati = {
            "percorso":    percorso_var.get().strip(),
            "risposta":    risposta_var.get().strip(),
            "mappa":       mappa_var.get().strip(),
            "date":        date_var.get(),
            "no_fallback": no_fallback_var.get(),
            "cifra":       cifra_var.get(),
            "revisiona":   revisiona_var.get(),
            # Se l'utente ha aperto la finestra associazioni, usa (e salva) le sue
            # righe; altrimenti None -> elabora riusa il file già salvato.
            "associazioni_righe": list(assoc_righe) if assoc_aperto[0] else None,
            "categorie_escluse": set(cat_escluse),
            "catasto_fittizio": catasto_fitt_var.get(),
        }
        # La cifratura chiede la password ORA, sul thread grafico principale.
        if scheda == 0 and dati["cifra"]:
            pw = _chiedi_password_gui()
            if pw is None:
                log("[INFO] Cifratura annullata.")
                bottone_esegui.configure(state="normal")
                return
            password_pre[0] = pw
        # Con la revisione servono dialoghi modali: si esegue sul thread
        # principale (i dialoghi tkinter non sono sicuri da un thread separato).
        if scheda == 0 and dati["revisiona"]:
            radice.after(50, lambda: _lavora(False, dati))
        else:
            threading.Thread(target=_lavora, args=(scheda == 1, dati), daemon=True).start()

    # --- layout ---
    note = ttk.Notebook(radice)
    tab_anon = ttk.Frame(note, padding=12)
    tab_rip  = ttk.Frame(note, padding=12)
    note.add(tab_anon, text="   Pseudonimizza   ")
    note.add(tab_rip,  text="   Ripristina   ")
    note.pack(fill="x", padx=10, pady=(10, 4))

    # scheda Pseudonimizza
    ttk.Label(tab_anon, text="Atto o cartella da pseudonimizzare:").grid(
        row=0, column=0, columnspan=3, sticky="w")
    ttk.Entry(tab_anon, textvariable=percorso_var).grid(
        row=1, column=0, columnspan=3, sticky="we", pady=(2, 6))
    b_file = ttk.Button(tab_anon, text="Sfoglia file…", command=scegli_file)
    b_file.grid(row=2, column=0, sticky="we", padx=(0, 4))
    b_cart = ttk.Button(tab_anon, text="Sfoglia cartella…", command=scegli_cartella)
    b_cart.grid(row=2, column=1, sticky="we", padx=4)
    b_apri = ttk.Button(tab_anon, text="Apri cartella", command=apri_cartella)
    b_apri.grid(row=2, column=2, sticky="we", padx=(4, 0))
    c_date = ttk.Checkbutton(tab_anon, text="Pseudonimizza anche le date generiche", variable=date_var)
    c_date.grid(row=3, column=0, columnspan=3, sticky="w", pady=(10, 0))
    c_cifra = ttk.Checkbutton(tab_anon, text="Cifra la mappa con una password", variable=cifra_var)
    c_cifra.grid(row=4, column=0, columnspan=3, sticky="w")
    c_fb = ttk.Checkbutton(tab_anon, text="Disattiva il riconoscimento dei nomi via regex (fallback)",
                           variable=no_fallback_var)
    c_fb.grid(row=5, column=0, columnspan=3, sticky="w")
    c_rev = ttk.Checkbutton(tab_anon, text="Rivedi a mano i termini dubbi (chiede conferma e ricorda le scelte)",
                            variable=revisiona_var)
    c_rev.grid(row=6, column=0, columnspan=3, sticky="w")
    c_catf = ttk.Checkbutton(tab_anon, text="Dati catastali: sostituisci con valori fittizi coerenti (invece del segnaposto)",
                             variable=catasto_fitt_var)
    c_catf.grid(row=8, column=0, columnspan=3, sticky="w")
    b_assoc = ttk.Button(tab_anon, text="Associazioni di nomi…", command=gestisci_associazioni)
    b_assoc.grid(row=7, column=0, sticky="w", pady=(8, 0))
    b_cat = ttk.Button(tab_anon, text="Categorie da non pseudonimizzare…", command=gestisci_categorie)
    b_cat.grid(row=7, column=1, columnspan=2, sticky="w", pady=(8, 0))
    for c in range(3):
        tab_anon.columnconfigure(c, weight=1)

    aggiungi_tooltip(b_file, "Scegli un singolo atto (.pdf, .txt o immagine) da pseudonimizzare.")
    aggiungi_tooltip(b_cart, "Scegli una cartella: verranno pseudonimizzati tutti gli atti al suo interno, con un'unica mappa.")
    aggiungi_tooltip(b_apri, "Apre nel sistema la cartella che contiene l'atto e i file prodotti.")
    aggiungi_tooltip(c_date, "Sostituisce anche le date generiche presenti nel testo, non solo quelle di nascita.")
    aggiungi_tooltip(c_cifra, "Protegge con una password il file della mappa (contiene i dati reali). Se attiva, la password verrà chiesta prima di elaborare.")
    aggiungi_tooltip(c_fb, "Disattiva il riconoscimento automatico dei nomi 'nudi' via regole: riduce i falsi positivi ma può lasciare qualche nome.")
    aggiungi_tooltip(c_rev, "Al termine chiede conferma parola per parola sui termini incerti; le tue scelte vengono ricordate.")
    aggiungi_tooltip(c_catf, "I numeri catastali (foglio, particella, sub) vengono sostituiti con numeri finti ma coerenti: lo stesso numero reale diventa sempre lo stesso finto, così i riferimenti restano leggibili. Reversibile con la mappa.")
    aggiungi_tooltip(b_assoc, "Raggruppa le diverse forme dello stesso soggetto (nome completo, abbreviazioni, «la Società») in un unico segnaposto. Vengono salvate accanto all'atto.")
    aggiungi_tooltip(b_cat, "Scegli quali tipi di dato LASCIARE nel testo (es. i dati catastali). Tutti gli altri vengono pseudonimizzati normalmente.")

    # scheda Ripristina
    ttk.Label(tab_rip, text="Risposta dell'AI (file con i segnaposto):").grid(
        row=0, column=0, columnspan=2, sticky="w")
    ttk.Entry(tab_rip, textvariable=risposta_var).grid(row=1, column=0, sticky="we", pady=(2, 6))
    b_srisp = ttk.Button(tab_rip, text="Sfoglia…", command=scegli_risposta)
    b_srisp.grid(row=1, column=1, padx=(6, 0))
    ttk.Label(tab_rip, text="Mappa privata (.json o .enc):").grid(
        row=2, column=0, columnspan=2, sticky="w")
    ttk.Entry(tab_rip, textvariable=mappa_var).grid(row=3, column=0, sticky="we", pady=(2, 6))
    b_smappa = ttk.Button(tab_rip, text="Sfoglia…", command=scegli_mappa)
    b_smappa.grid(row=3, column=1, padx=(6, 0))
    tab_rip.columnconfigure(0, weight=1)
    aggiungi_tooltip(b_srisp, "Scegli il file con la risposta dell'AI, quella che contiene i segnaposto tipo [PERSONA_1].")
    aggiungi_tooltip(b_smappa, "Scegli il file _mappa_privata.json (o .enc) creato durante la pseudonimizzazione.")

    # pulsante azione + log
    bottone_esegui = ttk.Button(radice, text="Elabora", command=on_esegui)
    bottone_esegui.pack(fill="x", padx=10, pady=6)
    aggiungi_tooltip(bottone_esegui, "Avvia l'operazione con le impostazioni scelte. L'esito compare nel riquadro sotto.")
    ttk.Label(radice, text="Esito:").pack(anchor="w", padx=12)
    area_log = scrolledtext.ScrolledText(radice, height=13, state="disabled", wrap="word")
    area_log.pack(fill="both", expand=True, padx=10, pady=(2, 4))

    # Disclaimer sempre visibile in fondo alla finestra.
    cornice_disc = tk.Frame(radice, bg="#fff3cd", bd=1, relief="solid")
    cornice_disc.pack(fill="x", padx=10, pady=(0, 10))
    tk.Label(cornice_disc, text="⚠  " + DISCLAIMER_TITOLO, bg="#fff3cd", fg="#7a5c00",
             font=("TkDefaultFont", 10, "bold"), anchor="w", justify="left").pack(fill="x", padx=8, pady=(6, 0))
    tk.Label(cornice_disc, text=DISCLAIMER_TESTO, bg="#fff3cd", fg="#5c4600",
             wraplength=700, justify="left", anchor="w").pack(fill="x", padx=8, pady=(0, 6))

    # Autore e crediti (in fondo alla finestra).
    tk.Label(radice, text=CREDITI_AUTORE, fg="#777777", font=("TkDefaultFont", 8)).pack(pady=(0, 0))
    tk.Label(radice, text=CREDITI_RINGRAZIAMENTI, fg="#777777", font=("TkDefaultFont", 8)).pack(pady=(0, 6))

    log("Pronto. Scegli un file o una cartella con «Sfoglia», poi premi «Elabora».")
    drena_coda()
    radice.mainloop()
    return 0
def main(argv=None) -> int:
    argv = sys.argv[1:] if argv is None else list(argv)
    # macOS può passare a un'app un argomento "-psn_..." (numero di serie del
    # processo): va ignorato, altrimenti argparse lo segnala come errore.
    argv = [a for a in argv if not a.startswith("-psn_")]

    ap = argparse.ArgumentParser(
        description="Pseudonimizzatore atti legali italiani (v7).",
        formatter_class=_FormatterIT, add_help=False)
    ap._positionals.title = "argomenti posizionali"
    ap._optionals.title = "opzioni"
    ap.add_argument("-h", "--help", "--aiuto", action="help", default=argparse.SUPPRESS,
                    help="Mostra questo messaggio di aiuto ed esci.")
    ap.add_argument("percorso", nargs="?", help="File .txt/.pdf/immagine o cartella.")
    ap.add_argument("--date", action="store_true", help="Pseudonimizza anche le date generiche.")
    ap.add_argument("--no-fallback-nomi", action="store_true", help="Disattiva il fallback regex per i nomi.")
    ap.add_argument("--encrypt-map", action="store_true", help="Cifra la mappa (pip install cryptography).")
    ap.add_argument("--ocr", action="store_true", help="Forza l'OCR su tutte le pagine dei PDF (per atti scansionati).")
    ap.add_argument("--revisione", action="store_true", help="Chiede conferma sui termini dubbi e ricorda le scelte.")
    ap.add_argument("--associa", action="store_true", help="Chiede le associazioni di nomi prima di elaborare (stesso segnaposto).")
    ap.add_argument("--escludi", metavar="CATEGORIE", help="Categorie da NON pseudonimizzare, separate da virgola (es. DATI_CATASTALI,TARGA).")
    ap.add_argument("--catasto-fittizio", action="store_true", help="Sostituisce i dati catastali con valori fittizi ma coerenti (invece del segnaposto).")
    ap.add_argument("--categorie", action="store_true", help="Elenca le categorie escludibili con --escludi ed esce.")
    ap.add_argument("--check", action="store_true", help="Mostra lo stato dei pacchetti installati ed esce.")
    ap.add_argument("--ripristina", nargs=2, metavar=("RISPOSTA_AI", "MAPPA"), help="Riapplica i dati reali.")
    ap.add_argument("--output", help="Output per --ripristina.")
    ap.add_argument("-i", "--interattivo", action="store_true", help="Avvia la modalità interattiva guidata (a testo).")
    ap.add_argument("-g", "--gui", action="store_true", help="Avvia l'interfaccia grafica (selezione file col mouse).")
    ap.add_argument("--no-gui", action="store_true", help="All'avvio senza argomenti, usa il menu a testo invece della grafica.")
    args = ap.parse_args(argv)

    try:
        if args.check:
            print("PSEUDONIMIZZATORE ATTI LEGALI v7 — stato dei pacchetti")
            stampa_stato_dipendenze()
            return 0
        if args.categorie:
            print("Categorie escludibili con --escludi (codice — descrizione):")
            for c, et in CATEGORIE_PSEUDONIMIZZABILI:
                print(f"  {c:<22} {et}")
            return 0
        categorie_escluse, _ign = parse_categorie(args.escludi)
        if _ign:
            print(f"[NOTA] Categorie non riconosciute ignorate: {', '.join(_ign)} "
                  f"(usa --categorie per l'elenco).", file=sys.stderr)
        # Interfaccia grafica: richiesta esplicita con --gui.
        if args.gui:
            return avvia_gui()
        # Nessun argomento oppure -i esplicito -> modalità guidata.
        # Se non è richiesto il testo (-i/--no-gui) e la grafica è disponibile,
        # si apre direttamente l'interfaccia grafica (es. con un doppio clic).
        if args.interattivo or not argv:
            if not args.interattivo and not args.no_gui and _gui_disponibile():
                return avvia_gui()
            codice = 0
            try:
                codice = modalita_interattiva()
            except KeyboardInterrupt:
                print("\n[ERRORE] Interrotto."); codice = 130
            except Exception as exc:
                print(f"[ERRORE] {exc}", file=sys.stderr); codice = 1
            _pausa_finale()
            return codice
        if args.ripristina:
            ripristina(args.ripristina[0], args.ripristina[1], args.output); return 0
        bersagli, cartella = risolvi_bersagli(args.percorso)
        associazioni_righe = None
        if args.associa:
            esistenti = leggi_associazioni_txt(percorso_associazioni(bersagli, cartella))
            associazioni_righe = modifica_associazioni_console(esistenti)
        elabora(bersagli, cartella, args.date, not args.no_fallback_nomi, args.encrypt_map, args.ocr,
                revisione=args.revisione,
                chiedi_dubbio=_chiedi_dubbio_console if args.revisione else None,
                associazioni_righe=associazioni_righe,
                categorie_escluse=categorie_escluse,
                catasto_fittizio=args.catasto_fittizio)
        return 0
    except KeyboardInterrupt:
        print("\n[ERRORE] Interrotto."); return 130
    except (ErroreUso, FileNotFoundError) as exc:
        print(f"[ERRORE] {exc}\n", file=sys.stderr)
        ap.print_help(sys.stderr)
        return 2
    except Exception as exc:
        print(f"[ERRORE] {exc}", file=sys.stderr); return 1

if __name__ == "__main__":
    raise SystemExit(main())
